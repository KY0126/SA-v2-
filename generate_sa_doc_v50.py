#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FJU Smart Hub v5.0 - System Analysis Document Generator
Generates a comprehensive Word document with auto TOC following the SA template format
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from datetime import datetime

# ─── Color Constants (Brand Colors) ───
BLUE = RGBColor(0, 49, 83)        # 聖母藍 #003153
GOLD = RGBColor(218, 165, 32)     # 梵蒂岡黃 #DAA520
GREEN = RGBColor(0, 128, 0)       # 嘉禾綠 #008000
RED = RGBColor(255, 0, 0)         # 警示紅 #FF0000
DARK_GRAY = RGBColor(51, 65, 85)  # #334155
LIGHT_GRAY = RGBColor(148, 163, 184)  # #94a3b8
WHITE = RGBColor(255, 255, 255)

def add_toc(doc):
    """Add automatic Table of Contents field"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)
    run._r.append(instrText)
    run._r.append(fld_char_end)

def set_cell_bg(cell, color_hex):
    """Set table cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=BLUE):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)
    return heading

def add_paragraph(doc, text, bold=False, size=10, color=DARK_GRAY, indent=0):
    """Add styled paragraph"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=1, size=10):
    """Add bullet point"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.left_indent = Cm(level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_header(table, headers, bg_hex='003153', text_color=WHITE):
    """Style table header row"""
    row = table.rows[0]
    for i, (cell, header) in enumerate(zip(row.cells, headers)):
        cell.text = header
        set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = text_color
                run.font.size = Pt(9)

def add_table_row(table, values, zebra=False):
    """Add data row to table"""
    row = table.add_row()
    bg = 'F8FAFC' if zebra else 'FFFFFF'
    for i, (cell, val) in enumerate(zip(row.cells, values)):
        cell.text = str(val)
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = DARK_GRAY
    return row

def add_mermaid_block(doc, title, code):
    """Add Mermaid diagram as code block with title"""
    p = doc.add_paragraph()
    run = p.add_run(f'【流程圖】{title}')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(4)

    # Code block style
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run(code)
    run2.font.name = 'Courier New'
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(30, 30, 30)
    # Set background for code block
    pPr = p2._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)

def generate_sa_document():
    doc = Document()

    # ──────────── Page Setup ────────────
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = '微軟正黑體'
    style.font.size = Pt(10)

    # ──────────── Cover Page ────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run('輔仁大學資訊管理學系')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('第三十三屆 系統分析規格書')
    run2.font.size = Pt(13)
    run2.font.color.rgb = DARK_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run('FJU Smart Hub')
    run_t.font.size = Pt(28)
    run_t.font.bold = True
    run_t.font.color.rgb = BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = subtitle.add_run('輔仁大學課外活動指導組\n全端校園智慧管理系統')
    run_s.font.size = Pt(14)
    run_s.font.color.rgb = GOLD
    run_s.font.bold = True

    doc.add_paragraph()

    # Gold divider line
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_d = divider.add_run('─' * 40)
    run_d.font.color.rgb = GOLD

    doc.add_paragraph()

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ('系統版本', 'v5.0'),
        ('技術堆疊', 'Laravel/PHP 8.3 · MySQL 8.0 · Vue 3 · Dify AI'),
        ('前端展示', 'Hono · Cloudflare Pages · Leaflet.js · GSAP'),
        ('文件版本', 'v5.0'),
        ('撰寫日期', datetime.now().strftime('%Y/%m/%d')),
        ('GitHub', 'https://github.com/KY0126/SA-v2-'),
    ]
    for i, (key, val) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        set_cell_bg(row.cells[0], '003153')
        set_cell_bg(row.cells[1], 'F8FAFC')
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ──────────── Modification Record ────────────
    add_heading(doc, '修改紀錄', 1)
    mod_table = doc.add_table(rows=1, cols=4)
    mod_table.style = 'Table Grid'
    add_table_header(mod_table, ['版本', '日期', '修改說明', '修改者'])
    add_table_row(mod_table, ['v1.0', '2026/01/10', '初稿建立，基礎架構設計', 'KY0126'], False)
    add_table_row(mod_table, ['v2.0', '2026/02/15', '新增 Dify AI 整合、三階段預約流程', 'KY0126'], True)
    add_table_row(mod_table, ['v3.0', '2026/03/01', '新增 E-Portfolio、器材盤點、信用系統', 'KY0126'], False)
    add_table_row(mod_table, ['v4.0', '2026/03/20', '完整 Bento Grid Dashboard、GSAP 動畫、已讀回條', 'KY0126'], True)
    add_table_row(mod_table, ['v5.0', '2026/04/03', '滾動式 Landing Page、ScrollTrigger、10大功能完整實作', 'KY0126'], False)
    doc.add_page_break()

    # ──────────── Table of Contents ────────────
    add_heading(doc, '目錄', 1)
    add_toc(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════════
    # CHAPTER 1: SYSTEM DESCRIPTION
    # ══════════════════════════════════════════════
    add_heading(doc, '第一章　系統描述', 1)

    # 1.1 需求分析與市場探索
    add_heading(doc, '一、需求分析與市場探索', 2)

    add_heading(doc, '問題陳述', 3)
    add_paragraph(doc, '''在輔仁大學的校園環境中，課外活動指導組（課指組）與學生社團之間存在多項根本性的管理痛點。首先，場地預約衝突問題頻繁發生：每當多個社團同時申請相同場地時，行政人員往往需要花費大量時間進行人工協調，既費時又缺乏系統性依據。其次，現行的活動申請流程高度依賴紙本作業，學生對相關法規（如《學生社團場地設備借用管理辦法》）的認識不足，導致退件率偏高，大幅降低行政效率。第三，器材借用管理缺乏電子化追蹤，逾期未還的情況無法自動提醒，造成器材遺失與帳目不清。最後，學生在多年校園生活中積累的社團參與、志工服務、幹部訓練等紀錄，缺乏系統化的整合與呈現平台，難以形成有說服力的個人學習履歷。''')

    add_heading(doc, '利害關係人分析', 3)
    stake_table = doc.add_table(rows=1, cols=3)
    stake_table.style = 'Table Grid'
    add_table_header(stake_table, ['利害關係人', '角色描述', '主要需求'])
    stakeholders = [
        ('一般學生', '校內在學生，使用系統預約場地、借用器材、查詢社團資訊', '便捷申請流程、即時狀態查詢、個人履歷管理、信用點數追蹤'),
        ('社團幹部', '各社團正式幹部，負責社團活動規劃與申請', 'AI 企劃生成、協商機制、幹部證書自動化、社員管理'),
        ('指導教授', '社團指導老師，需掌握社團運作狀況', '學生職能成長追蹤、績效雷達圖、活動熱力圖'),
        ('課指組/行政', '課外活動指導組行政人員，審核申請、管理場地', 'AI 智慧預審、統計報表、SDGs 貢獻分析、審核流程自動化'),
        ('資訊中心/IT', '系統管理員，負責技術維護與安全', 'WAF 日誌、API 監控、R2 存儲管理、2FA 設定'),
    ]
    for i, (role, desc, need) in enumerate(stakeholders):
        add_table_row(stake_table, [role, desc, need], i % 2 == 1)

    doc.add_paragraph()
    add_heading(doc, '需求蒐集方法', 3)
    add_paragraph(doc, '本系統透過以下多元方法蒐集需求：')
    methods = [
        '深度訪談：與課外活動指導組行政人員進行 3 次深度訪談（每次 60-90 分鐘），梳理現行流程痛點',
        '使用者問卷：對 87 個社團幹部發放線上問卷（回收率 73%），統計最常遇到的問題類型',
        '現場觀察：觀察 2 學期（上下學期各 4 週）的場地申請與器材借用作業流程',
        '競爭者分析：研究 HiTeach、Slido、飛鳶平台等系統的功能設計',
        '法規解析：深度索引輔仁大學課外活動官方網站（5 個主要法規頁面）建立 Knowledge Base',
    ]
    for m in methods:
        add_bullet(doc, m)

    add_heading(doc, '競爭者分析', 3)
    comp_table = doc.add_table(rows=1, cols=5)
    comp_table.style = 'Table Grid'
    add_table_header(comp_table, ['系統', 'AI 預審', '場地協商', '信用系統', '多語言'])
    competitors = [
        ('FJU Smart Hub', '✅ RAG+Dify', '✅ 3/6分鐘規則', '✅ 60分停權', '✅ 5國語言'),
        ('HiTeach', '❌', '❌', '❌', '✅'),
        ('飛鳶平台', '部分', '❌', '❌', '❌'),
        ('Slido', '❌', '❌', '❌', '✅'),
        ('傳統紙本', '❌', '❌', '❌', '❌'),
    ]
    for i, row_data in enumerate(competitors):
        add_table_row(comp_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()
    add_heading(doc, '市場定位/利基', 3)
    add_paragraph(doc, '現有的課程輔助系統（如 HiTeach、Slido）主要聚焦於課堂即時互動，並未針對「校園資源管理」與「社團行政效率」設計。FJU Smart Hub 定位於：（1）整合 AI 技術自動化行政流程；（2）提供場地與器材的智慧化協商機制；（3）構建學生個人學習歷程平台。此定位在國內大學系統中具有高度創新性與實用價值。')

    # 1.2 系統發展目的
    add_heading(doc, '二、系統發展目的', 2)
    add_paragraph(doc, '本系統旨在打造一個「智慧、高效、透明」的校園資源管理生態系統，具體目的如下：')

    purposes = [
        ('對學生', '提供一站式校園服務入口，包含場地預約、器材借用、社團查詢、活動參加及個人職能履歷管理，顯著降低行政往返成本'),
        ('對社團幹部', '透過 AI 企劃生成器、幹部證書自動化及傳承檢查表，減少重複性行政工作，讓幹部專注於活動品質提升'),
        ('對指導教授', '提供量化的社團績效數據（雷達圖、熱力圖、職能 Spider 圖），協助教授掌握社員成長軌跡'),
        ('對行政人員', '透過 AI 智慧預審（RAG）自動過濾 80% 的問題申請，並提供即時統計報表與 SDGs 貢獻分析，支援決策制定'),
        ('對資訊中心', '提供完整的系統監控儀表板，包含 API 延遲、WAF 日誌、R2 使用率，確保系統安全與穩定運行'),
    ]

    for target, desc in purposes:
        p = doc.add_paragraph()
        run1 = p.add_run(f'◆ {target}：')
        run1.font.bold = True
        run1.font.color.rgb = BLUE
        run1.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.3)

    add_heading(doc, '功能性需求', 3)
    func_reqs = [
        'FR-01：使用者身份驗證（Google OAuth + hd 域名檢查 + 2FA TOTP/SMS）',
        'FR-02：場地查詢、預約申請、三階段流程（志願序→協商→核定）',
        'FR-03：3/6 分鐘精確協商機制（GPT-4 介入 + 紅光扣分動畫）',
        'FR-04：AI 智慧預審（Dify RAG → risk_level Medium/High/Low 輸出）',
        'FR-05：器材庫存管理、借用申請、LINE/SMS 到期提醒、QR Code 取件',
        'FR-06：E-Portfolio 職能履歷、技能標籤、PDF 匯出（outbox 異步）',
        'FR-07：AI 自動企劃生成器（呼叫 Dify Workflow + 法規合規性檢查）',
        'FR-08：幹部證書自動化產出（含數位簽章）',
        'FR-09：信用積分系統（觀察者模式監控 < 60 分強制 JWT 失效）',
        'FR-10：已讀回條追蹤（Unique Token + 10 秒強制閱讀 Modal）',
        'FR-11：校園互動地圖（Leaflet.js + SVG Overlay + flyTo 動畫）',
        'FR-12：玻璃行事曆（GSAP 右滑 + 衝突高亮 + 協商按鈕）',
        'FR-13：動態活動牆（標籤/關鍵字搜尋 + 分類卡片）',
        'FR-14：數位時光膠囊（R2 文件封裝移交）',
        'FR-15：角色專屬儀表板（5 種角色各有不同圖表組合）',
        'FR-16：五國語言 i18n（繁中/簡中/英/日/韓，全系統即時切換）',
        'FR-17：多通路通知鏈（LINE Notify + SMS + SMTP + Outlook）',
        'FR-18：PWA 離線支援（確保地下室環境可讀取憑證）',
    ]
    for req in func_reqs:
        add_bullet(doc, req)

    add_heading(doc, '非功能性需求', 3)
    nonfunc_reqs = [
        'NFR-01 效能：場地查詢 API < 200ms，地圖渲染 < 3s，整體 API 可用率 > 99.5%',
        'NFR-02 安全：Cloudflare WAF + Turnstile 人機驗證、JWT HttpOnly Cookie、防抵賴 Unique Token',
        'NFR-03 穩定性：Redis 快取 2FA 與任務計時、MySQL 悲觀鎖防止 Race Condition',
        'NFR-04 可維護性：Docker 容器化部署、環境變數管理（禁止 API Key 寫入程式碼）',
        'NFR-05 無障礙性：色彩對比符合 WCAG 2.1 AA 標準，無障礙設施地圖標記',
        'NFR-06 可擴展性：Cloudflare Workers 邊緣部署，支援全球 CDN 加速',
    ]
    for req in nonfunc_reqs:
        add_bullet(doc, req)

    # 1.3 系統範圍
    add_heading(doc, '三、系統範圍', 2)
    add_paragraph(doc, '本系統涵蓋輔仁大學課外活動指導組之全部數位化業務，主要功能模組範圍如下：')

    scope_items = [
        ('校園地圖模組', '互動式 Leaflet.js 地圖，標示所有建築、場地狀態、無障礙設施'),
        ('場地預約模組', '三階段預約流程（志願序→AI 預審→衝突協商→官方核定）'),
        ('器材借用模組', '庫存管理、借用申請、QR Code 取件、逾期追蹤'),
        ('社團管理模組', 'CRUD 操作、幹部管理、社員名冊、傳承機制'),
        ('活動管理模組', '活動牆、日曆管理、AI 企劃生成、幹部證書'),
        ('E-Portfolio 模組', '職能標籤、活動紀錄、PDF 匯出、雷達圖分析'),
        ('AI 核心模組', 'RAG 法規查詢、智慧預審、申訴摘要、AI 導覽'),
        ('行政報表模組', 'SDGs 貢獻、場地使用率、社團評鑑統計'),
        ('安全模組', '2FA 驗證、信用積分、已讀回條、防抵賴追蹤'),
        ('系統管理模組', '用戶管理、WAF 日誌、R2 儲存監控'),
    ]
    scope_table = doc.add_table(rows=1, cols=2)
    scope_table.style = 'Table Grid'
    add_table_header(scope_table, ['功能模組', '範圍說明'])
    for i, (mod, desc) in enumerate(scope_items):
        add_table_row(scope_table, [mod, desc], i % 2 == 1)

    doc.add_paragraph()

    # 1.4 背景知識
    add_heading(doc, '四、背景知識', 2)
    bg_topics = [
        ('RAG（Retrieval-Augmented Generation）', '結合檢索與生成的 AI 架構。本系統透過 Dify 平台建立輔仁大學法規 Knowledge Base，當使用者提交預約申請時，系統自動檢索相關法規條文，並由 LLM（大型語言模型）生成合規性分析報告，標記 Low/Medium/High 風險等級。'),
        ('信用積分制（Credit Score System）', '參考 Uber 司機評分與芝麻信用的設計理念，本系統對每位用戶維護 0-100 分的信用積分。正向行為（準時簽到、協商成功）加分，負向行為（無故缺席、協商逾期）扣分。低於 60 分觸發強制登出機制，確保校園資源使用的公正性。'),
        ('JWT HttpOnly Cookie 安全機制', 'JSON Web Token 存儲於 HttpOnly Cookie 中，防止 XSS 攻擊竊取 Token。搭配 Redis 管理 Token 黑名單，實現即時 JWT 失效（信用分數 < 60 時由觀察者模式觸發）。'),
        ('Cloudflare WAF + Turnstile', 'Cloudflare Web Application Firewall 提供第一層防護，針對 @cloud.fju.edu.tw 帳號設置 Skip Rules。Turnstile 作為非互動式人機驗證，替代傳統 CAPTCHA，在預約/申復流程中防止自動化攻擊。'),
        ('WebSocket 即時通訊', '協商對話室採用 Socket.io 實現雙向即時通訊。3/6 分鐘計時器在 Redis 中維護，確保分散式環境下的計時精確性。'),
    ]
    for title, content in bg_topics:
        p = doc.add_paragraph()
        run1 = p.add_run(f'■ {title}：')
        run1.font.bold = True
        run1.font.size = Pt(10)
        run1.font.color.rgb = BLUE
        p.paragraph_format.space_after = Pt(2)

        p2 = doc.add_paragraph()
        run2 = p2.add_run(content)
        run2.font.size = Pt(9.5)
        run2.font.color.rgb = DARK_GRAY
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(10)

    # 1.5 系統限制
    add_heading(doc, '五、系統限制', 2)
    limits = [
        '帳號限制：僅接受 @cloud.fju.edu.tw 及 @fju.edu.tw 網域信箱，外部帳號無法使用',
        '場地資料：場地資訊（建築位置、容量、無障礙設施）需由行政人員定期維護更新',
        'AI 精準度：Dify RAG 預審結果為輔助性建議，最終審核仍需人工確認',
        '網路需求：互動地圖（Leaflet.js）需要穩定網路連線，地下室環境透過 PWA 離線模式提供基本功能',
        'Dify API 限制：AI 服務依賴外部 Dify 平台，受 API 速率限制（Rate Limit）影響',
        '瀏覽器相容性：建議使用 Chrome 90+、Firefox 88+、Safari 14+，不支援 IE 11 及以下版本',
    ]
    for limit in limits:
        add_bullet(doc, limit)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # CHAPTER 2: SOFTWARE REQUIREMENT SPEC
    # ══════════════════════════════════════════════
    add_heading(doc, '第二章　軟體需求規格', 1)

    # 2.1 功能需求
    add_heading(doc, '一、功能需求', 2)

    add_heading(doc, '使用者角色說明表', 3)
    role_table = doc.add_table(rows=1, cols=4)
    role_table.style = 'Table Grid'
    add_table_header(role_table, ['角色代碼', '角色名稱', '說明', '主要功能權限'])
    roles = [
        ('STUDENT', '一般學生', '輔仁大學在學學生，使用 @cloud.fju.edu.tw 帳號', '場地預約、器材借用、活動參加、E-Portfolio 管理、AI 查詢'),
        ('CLUB_OFFICER', '社團幹部', '經核准之社團幹部，擁有社團管理權限', '活動申請、AI 企劃生成、幹部證書、社員管理、衝突協商'),
        ('PROFESSOR', '指導教授', '社團指導教授，可查看所指導社團之所有資訊', '績效雷達查看、活動熱力圖、職能追蹤、評鑑報告'),
        ('ADMIN', '行政人員', '課指組行政人員（處室），擁有審核與管理權限', '申請審核、AI 預審結果查閱、統計報表、SDGs 分析'),
        ('IT_ADMIN', '系統管理員', '資訊中心技術人員，負責系統維護', '用戶管理、WAF 日誌、API 監控、R2 存儲管理、2FA 設定'),
    ]
    for i, row_data in enumerate(roles):
        add_table_row(role_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()

    add_heading(doc, '使用者故事對應表', 3)
    story_table = doc.add_table(rows=1, cols=5)
    story_table.style = 'Table Grid'
    add_table_header(story_table, ['Epic', '使用者故事 ID', '使用者故事', '優先級', '對應角色'])

    stories = [
        ('場地管理', 'US-01', '身為社團幹部，我希望能線上申請場地並知道 AI 預審結果，以便提早知道是否需要補件', 'High', '社團幹部'),
        ('場地管理', 'US-02', '身為一般學生，我希望透過地圖快速查詢場地的無障礙設施，以確保活動無障礙', 'High', '一般學生'),
        ('場地管理', 'US-03', '身為行政人員，我希望 AI 自動標記高風險申請，以便優先處理需人工審核的案件', 'High', '行政人員'),
        ('場地管理', 'US-04', '身為社團幹部，我希望場地衝突時能線上協商，在 6 分鐘內解決，以避免申請延誤', 'High', '社團幹部'),
        ('器材借用', 'US-05', '身為一般學生，我希望線上申請借用器材並掃 QR Code 取件，以節省往返辦公室的時間', 'High', '一般學生'),
        ('器材借用', 'US-06', '身為行政人員，我希望系統自動發送 LINE 提醒器材到期，以減少追催工作', 'Medium', '行政人員'),
        ('E-Portfolio', 'US-07', '身為一般學生，我希望系統自動彙整我的活動紀錄並生成 PDF 履歷，以備求職使用', 'High', '一般學生'),
        ('E-Portfolio', 'US-08', '身為指導教授，我希望查看職能雷達圖追蹤學生成長，以提供精準指導', 'Medium', '指導教授'),
        ('AI 功能', 'US-09', '身為社團幹部，我希望 AI 自動生成符合法規的活動企劃書，以節省撰寫時間', 'High', '社團幹部'),
        ('AI 功能', 'US-10', '身為一般學生，我希望向 AI 查詢場地借用法規，以避免因不了解規定而申請失敗', 'Medium', '一般學生'),
        ('信用系統', 'US-11', '身為行政人員，我希望系統自動追蹤用戶違規並扣分，並在 60 分以下強制停權', 'High', '行政人員'),
        ('通知系統', 'US-12', '身為一般學生，我希望重要通知強制彈出 10 秒閱讀視窗，且系統記錄我的閱讀時間', 'High', '一般學生'),
        ('多語言', 'US-13', '身為外籍交換生，我希望整個系統能切換為英文/日文/韓文，以無障礙使用系統', 'Medium', '全部角色'),
    ]
    for i, row_data in enumerate(stories):
        add_table_row(story_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()

    # User Story Cards
    add_heading(doc, '使用者故事卡（核心功能）', 3)

    story_cards = [
        {
            'id': 'US-01', 'role': '社團幹部', 'epic': '場地管理',
            'story': '身為社團幹部，我希望能線上提交場地預約申請，並在送出後立即收到 AI 智慧預審結果，以便在核定前預先了解是否需要補充文件。',
            'value': '減少退件次數，提升申請成功率，縮短行政流程',
            'acceptance': [
                '系統驗證使用者為 @cloud.fju.edu.tw 帳號且信用分數 ≥ 60 分',
                '申請表單包含：場地選擇、活動日期/時段（志願序 1-3）、活動說明（500字以內）、預計人數',
                'AI 預審在 3 秒內回傳 risk_level（Low/Medium/High）及 reasoning 說明',
                'High 風險申請自動拒絕並附上法規說明；Medium 轉人工審核；Low 自動核准',
                '申請成功後，系統透過 Outlook（Microsoft Graph API）發送確認信',
            ]
        },
        {
            'id': 'US-04', 'role': '社團幹部', 'epic': '衝突協商',
            'story': '身為申請方社團幹部，當場地預約發生衝突時，我希望能進入即時協商室與對方溝通，並在 GPT-4 AI 的協助下，於最短時間內達成共識。',
            'value': '取代人工電話協調，提高協商效率，確保程序透明',
            'acceptance': [
                'LINE Notify 自動向雙方發送協商邀請連結及對方聯絡資訊',
                '協商室即時顯示 WebSocket 對話，內建語意過濾（偵測與協商無關的訊息）',
                '3 分鐘靜默計時（Redis）：雙方無訊息時啟動，到時 GPT-4 介入提供 3 個具體建議',
                '6 分鐘強制結束：若仍無共識，系統自動扣除雙方各 10 信用點並記錄 credit_logs',
                '協商成功：雙方點擊「完成協商」後，更新 conflicts 表狀態，重導至行事曆管理頁面',
            ]
        },
    ]

    for card in story_cards:
        card_table = doc.add_table(rows=5, cols=2)
        card_table.style = 'Table Grid'
        card_labels = ['故事 ID', '角色', 'Epic', '故事描述', '商業價值']
        card_values = [card['id'], card['role'], card['epic'], card['story'], card['value']]
        for i, (label, val) in enumerate(zip(card_labels, card_values)):
            row = card_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = val
            set_cell_bg(row.cells[0], '003153')
            set_cell_bg(row.cells[1], 'FFFFFF' if i % 2 == 0 else 'F8FAFC')
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = WHITE
                    run.font.size = Pt(9)
            for para in row.cells[1].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

        p_ac = doc.add_paragraph()
        run_ac = p_ac.add_run(f'驗收標準（Acceptance Criteria）— {card["id"]}：')
        run_ac.font.bold = True
        run_ac.font.size = Pt(9.5)
        run_ac.font.color.rgb = BLUE
        p_ac.paragraph_format.space_before = Pt(8)
        for criterion in card['acceptance']:
            add_bullet(doc, criterion, size=9)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # CHAPTER 3: SOFTWARE DESIGN SPEC
    # ══════════════════════════════════════════════
    add_heading(doc, '第三章　軟體設計規格', 1)

    # 3.1 Database Design
    add_heading(doc, '一、資料庫設計', 2)

    add_heading(doc, '關連一覽表', 3)
    rel_table = doc.add_table(rows=1, cols=4)
    rel_table.style = 'Table Grid'
    add_table_header(rel_table, ['資料表名稱', '中文說明', '主鍵', '外鍵關聯'])
    db_tables = [
        ('users', '使用者帳號表', 'id', '無'),
        ('venues', '場地資訊表', 'id', '無'),
        ('equipment', '器材庫存表', 'id', '無'),
        ('reservations', '場地預約記錄', 'id', 'user_id → users, venue_id → venues'),
        ('equipment_loans', '器材借用記錄', 'id', 'user_id → users, equipment_id → equipment'),
        ('clubs', '社團資訊表', 'id', 'advisor_id → users'),
        ('club_members', '社團成員對應', 'id', 'user_id → users, club_id → clubs'),
        ('activities', '活動記錄表', 'id', 'club_id → clubs, venue_id → venues'),
        ('conflicts', '場地衝突記錄', 'id', 'reservation_id_a, reservation_id_b → reservations'),
        ('credit_logs', '信用點數日誌', 'id', 'user_id → users'),
        ('notification_logs', '通知已讀日誌', 'id', 'user_id → users'),
        ('certificates', '幹部證書記錄', 'id', 'user_id → users, club_id → clubs'),
        ('portfolios', 'E-Portfolio 記錄', 'id', 'user_id → users'),
        ('ai_screening_logs', 'AI 預審記錄', 'id', 'reservation_id → reservations'),
        ('map_elements', '地圖建築元素', 'id', '無'),
    ]
    for i, row_data in enumerate(db_tables):
        add_table_row(rel_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()
    add_heading(doc, '關連結構表（核心資料表）', 3)

    # Users table
    add_paragraph(doc, '▶ users（使用者帳號表）', bold=True, color=BLUE)
    users_table = doc.add_table(rows=1, cols=5)
    users_table.style = 'Table Grid'
    add_table_header(users_table, ['欄位名稱', '資料型態', '說明', '約束條件', '備註'])
    users_fields = [
        ('id', 'INT', '使用者唯一識別碼', 'PK, AI, NN', ''),
        ('name', 'VARCHAR(50)', '真實姓名', 'NN', ''),
        ('student_id', 'VARCHAR(20)', '學號/工號', 'UNIQUE, NN', '111071xxx'),
        ('email', 'VARCHAR(100)', '校園 Email', 'UNIQUE, NN', '@cloud.fju.edu.tw 或 @fju.edu.tw'),
        ('phone', 'VARCHAR(20)', '手機號碼', 'NULL', '用於 SMS 2FA'),
        ('role', "ENUM('student','club_officer','professor','admin','it_admin')", '使用者角色', 'NN', 'DEFAULT student'),
        ('credit_score', 'TINYINT UNSIGNED', '信用積分 0-100', 'NN', 'DEFAULT 100, 低於60強制停權'),
        ('google_id', 'VARCHAR(100)', 'Google OAuth ID', 'UNIQUE', 'hd 域名驗證後存入'),
        ('totp_secret', 'VARCHAR(100)', 'TOTP 2FA 金鑰', 'NULL', 'PHPGangsta 加密儲存'),
        ('is_active', 'TINYINT(1)', '帳號是否啟用', 'NN', 'DEFAULT 1'),
        ('last_login_at', 'DATETIME', '最後登入時間', 'NULL', ''),
        ('created_at', 'DATETIME', '建立時間', 'NN', 'DEFAULT CURRENT_TIMESTAMP'),
        ('updated_at', 'DATETIME', '更新時間', 'NN', 'ON UPDATE CURRENT_TIMESTAMP'),
    ]
    for i, row_data in enumerate(users_fields):
        add_table_row(users_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()

    # Reservations table
    add_paragraph(doc, '▶ reservations（場地預約記錄表）', bold=True, color=BLUE)
    res_table = doc.add_table(rows=1, cols=5)
    res_table.style = 'Table Grid'
    add_table_header(res_table, ['欄位名稱', '資料型態', '說明', '約束條件', '備註'])
    res_fields = [
        ('id', 'INT', '預約唯一識別碼', 'PK, AI, NN', ''),
        ('user_id', 'INT', '申請者 ID', 'FK → users, NN', ''),
        ('venue_id', 'INT', '場地 ID', 'FK → venues, NN', ''),
        ('start_time', 'DATETIME', '活動開始時間', 'NN', ''),
        ('end_time', 'DATETIME', '活動結束時間', 'NN', ''),
        ('purpose', 'TEXT', '活動說明', 'NN', 'AI 預審用'),
        ('expected_people', 'SMALLINT', '預計人數', 'NN', ''),
        ('weight_level', 'TINYINT', '權限等級', 'NN', '1=校方/2=學生自治/3=一般社團'),
        ('priority_slot', 'TINYINT', '志願序順序', 'NN', '1-3'),
        ('status', "ENUM('PENDING','PENDING_MANUAL_REVIEW','APPROVED','REJECTED','CONFLICT','CANCELLED')", '申請狀態', 'NN', 'DEFAULT PENDING'),
        ('ai_risk_level', "ENUM('Low','Medium','High')", 'AI 預審風險等級', 'NULL', ''),
        ('ai_reasoning', 'TEXT', 'AI 預審推理說明', 'NULL', ''),
        ('pdf_url', 'VARCHAR(500)', 'PDF 申請單 URL', 'NULL', 'Cloudflare R2'),
        ('totp_qr_code', 'VARCHAR(500)', 'TOTP QR Code URL', 'NULL', ''),
        ('approved_by', 'INT', '審核者 ID', 'FK → users, NULL', ''),
        ('approved_at', 'DATETIME', '審核時間', 'NULL', ''),
        ('created_at', 'DATETIME', '建立時間', 'NN', 'DEFAULT CURRENT_TIMESTAMP'),
    ]
    for i, row_data in enumerate(res_fields):
        add_table_row(res_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()

    # Credit logs table
    add_paragraph(doc, '▶ credit_logs（信用點數日誌表）', bold=True, color=BLUE)
    credit_table = doc.add_table(rows=1, cols=5)
    credit_table.style = 'Table Grid'
    add_table_header(credit_table, ['欄位名稱', '資料型態', '說明', '約束條件', '備註'])
    credit_fields = [
        ('id', 'INT', '日誌唯一識別碼', 'PK, AI, NN', ''),
        ('user_id', 'INT', '用戶 ID', 'FK → users, NN', ''),
        ('action', 'VARCHAR(100)', '動作類型', 'NN', '如：準時簽到、協商逾期'),
        ('change_amount', 'TINYINT', '點數變更量', 'NN', '正值加分，負值扣分'),
        ('score_after', 'TINYINT UNSIGNED', '變更後積分', 'NN', ''),
        ('reason', 'TEXT', '原因說明', 'NULL', ''),
        ('triggered_by', 'VARCHAR(50)', '觸發來源', 'NULL', '如：OBSERVER_PATTERN, MANUAL'),
        ('created_at', 'DATETIME', '記錄時間', 'NN', 'DEFAULT CURRENT_TIMESTAMP'),
    ]
    for i, row_data in enumerate(credit_fields):
        add_table_row(credit_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()

    # ER Diagram description
    add_heading(doc, '資料庫 E-R 圖說明', 3)
    add_paragraph(doc, '系統資料庫採用 MySQL 8.0 InnoDB 引擎，所有外鍵關聯均使用 ON DELETE CASCADE 確保資料一致性。主要實體關聯如下：')

    er_items = [
        'users (1) ──── (N) reservations：一名用戶可以建立多筆預約記錄',
        'users (1) ──── (N) equipment_loans：一名用戶可以有多筆器材借用記錄',
        'users (N) ────(M) clubs（透過 club_members）：一名用戶可加入多個社團，一個社團有多名成員',
        'venues (1) ──── (N) reservations：一個場地可以被多次預約',
        'reservations (1) ──── (1) ai_screening_logs：每筆預約對應一份 AI 預審記錄',
        'reservations (1) ──── (1) conflicts（選擇性）：衝突時產生對應的協商記錄',
        'users (1) ──── (N) credit_logs：用戶的每次信用分數變動都記錄在 credit_logs',
        'clubs (1) ──── (N) activities：一個社團可以舉辦多個活動',
        'users (1) ──── (1) portfolios：每位用戶對應一份 E-Portfolio（一對一關係）',
    ]
    for item in er_items:
        add_bullet(doc, item, size=9)

    doc.add_paragraph()

    # 3.2 Interface Design
    add_heading(doc, '二、介面設計', 2)

    add_heading(doc, '介面藍圖一覽表', 3)
    ui_table = doc.add_table(rows=1, cols=4)
    ui_table.style = 'Table Grid'
    add_table_header(ui_table, ['介面編號', '介面名稱', '使用角色', '主要功能'])
    ui_pages = [
        ('UI-01', 'Landing Page（動態介紹頁）', '訪客', '活動輪播、痛點說明、功能 Bento Grid、圓形評價輪播、即時快訊橫幅'),
        ('UI-02', 'Portal 登入頁', '訪客', 'Glassmorphism 登入盒、角色選擇、Google OAuth、Cloudflare Turnstile'),
        ('UI-03', 'Dashboard 主頁（Bento Grid）', '全部', '互動地圖（底層）+ 玻璃行事曆（頂層）+ 側邊欄'),
        ('UI-04', '場地預約頁', '學生/幹部', '三階段進度指示、場地選擇、AI 預審結果卡片'),
        ('UI-05', '衝突協商室', '學生/幹部', 'WebSocket 聊天、計時器、AI 建議面板、紅光動畫'),
        ('UI-06', '器材借用頁', '學生/幹部', '庫存表格、借用申請、QR Code Modal'),
        ('UI-07', '社團管理頁', '幹部/行政', 'CRUD 操作、社員名冊、統計圖表'),
        ('UI-08', '動態活動牆', '全部', '卡片 Grid、標籤篩選、關鍵字搜尋'),
        ('UI-09', 'E-Portfolio 頁', '學生/教授', '個人資訊卡、活動紀錄、職能雷達圖、PDF 下載'),
        ('UI-10', 'AI 工具頁', '全部', 'Tab：AI 導覽 / 法規 RAG / 預約流程 RAG / 企劃生成器'),
        ('UI-11', '角色儀表板', '全部', '各角色專屬圖表（折線/圓餅/雷達/Gauge/熱力圖）'),
        ('UI-12', '行事曆管理頁', '全部', '全月視圖、衝突高亮、協商入口'),
        ('UI-13', '用戶管理頁', '行政/IT', '用戶 CRUD、信用日誌查看、角色管理'),
        ('UI-14', '系統設定頁', '全部', '帳號設定、2FA 管理、通知偏好、已讀回條查詢'),
    ]
    for i, row_data in enumerate(ui_pages):
        add_table_row(ui_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()
    add_heading(doc, '介面藍圖：Dashboard 主頁（UI-03）', 3)
    add_paragraph(doc, '設計概念：雙模態交互介面（Dual-Modal Interface）')
    add_paragraph(doc, '''
主頁採用創新的雙層疊加設計，底層為全螢幕 Leaflet.js 互動地圖，頂層為右側浮動玻璃行事曆面板：

• 頁首（Header）：左側輔大校徽 + 「FJU Smart Hub」文字標誌；中間懸浮搜尋列（仿 Google Maps 設計，圓角、半透明背景）；右側含通知紅點鈴鐺與學號首字母頭像。

• 左側側邊欄（30% 寬度）：垂直排列四個功能分組：
  - [主要功能] 儀表板、校園地圖、場地/設備借用、行事曆
  - [社群與社團] 社團資訊、活動牆
  - [AI 核心專區] AI 資訊/導覽、法規查詢 RAG、場地與器材預約流程 RAG
  - [管理與報表] 報修、申訴、統計

• 側邊欄頂部：個人信用儀表板（黃色 Gauge 圓弧 + 健康狀態標籤）

• 右側主視覺（70%）：全螢幕 Leaflet.js 地圖，建築標記點擊後觸發 AJAX 獲取即時維修狀態，並以動畫展示建築資訊卡。

• 玻璃行事曆面板：backdrop-filter: blur(15px)，停靠右緣 40% 寬度，GSAP 右向左滑入，顯示預約狀況，若時段衝突提供「協商」按鈕。

• 浮動 FAB（右下角）：動態柴犬 AI 圖標，提供快速 AI 查詢入口。
    '''.strip())

    add_heading(doc, '介面藍圖：衝突協商室（UI-05）', 3)
    add_paragraph(doc, '''
衝突協商室採用全螢幕 Modal 設計：
• 頂部：場地名稱、衝突時段、對方社團資訊
• 中央：WebSocket 即時對話框（區分己方訊息/對方訊息/AI訊息/系統訊息）
• 計時器：醒目位置顯示 MM:SS 倒計時，接近 3 分鐘時變紅閃爍
• 底部：輸入框 + 發送按鈕
• AI 建議面板：3 分鐘後展開，顯示 GPT-4 提供的 3 個具體解決方案按鈕
• 6 分鐘觸發：全螢幕紅色脈衝動畫覆蓋，並顯示扣分通知
    '''.strip())

    # 3.3 API Definitions
    add_heading(doc, '三、完整 API 接口定義', 2)

    api_table = doc.add_table(rows=1, cols=5)
    api_table.style = 'Table Grid'
    add_table_header(api_table, ['HTTP 方法', 'API 路徑', '功能描述', '請求參數', '回應格式'])
    apis = [
        ('GET', '/api/auth/me', '取得目前登入用戶資訊', 'Authorization Header（JWT）', '{ user_id, name, role, credit_score }'),
        ('POST', '/api/auth/google', 'Google OAuth 登入（hd 域名驗證）', '{ code: string }', '{ token, user }'),
        ('POST', '/api/auth/2fa/verify', '驗證 TOTP/SMS 2FA 代碼', '{ user_id, code, type }', '{ success, token }'),
        ('GET', '/api/venues', '取得所有場地列表', 'Query: type, accessible', '[ { id, name, capacity, accessible } ]'),
        ('POST', '/api/reservations', '提交場地預約申請（含 AI 預審）', '{ venue_id, start_time, end_time, purpose, weight_level }', '{ id, status, ai_screening: { risk_level, reasoning } }'),
        ('GET', '/api/reservations', '取得我的預約記錄', 'Query: status, page', '[ { id, venue, status, ai_risk_level } ]'),
        ('PATCH', '/api/reservations/:id/status', '更新預約狀態（行政審核）', '{ status, reason }', '{ id, status, updated_at }'),
        ('GET', '/api/conflicts', '取得衝突列表', 'Query: status', '[ { id, venue, party1, party2, status } ]'),
        ('PATCH', '/api/conflicts/:id/resolve', '解決場地衝突', '{ resolution, winner_party }', '{ id, status: resolved }'),
        ('POST', '/api/conflicts/:id/penalty', '觸發協商逾期扣分（6分鐘機制）', '{ user_id, reason }', '{ conflict_id, penalty: -10 }'),
        ('GET', '/api/equipment', '取得器材庫存列表', 'Query: category', '[ { id, name, qty_available } ]'),
        ('POST', '/api/equipment/borrow', '提交器材借用申請', '{ equipment_id, return_date }', '{ id, status, totp_qr }'),
        ('GET', '/api/clubs', '取得社團列表', 'Query: category, keyword', '[ { id, name, members, credit } ]'),
        ('GET', '/api/clubs/:id', '取得社團詳情', 'Path: id', '{ id, name, advisor, activities }'),
        ('GET', '/api/activities', '取得活動牆列表', 'Query: keyword, tag', '[ { id, title, club, tags } ]'),
        ('GET', '/api/dashboard/:role', '取得角色專屬儀表板數據', 'Path: role', '{ stats, charts[] }'),
        ('GET', '/api/portfolio/:userId', '取得用戶 E-Portfolio', 'Path: userId', '{ skills, activities, certificates }'),
        ('GET', '/api/calendar', '取得月曆活動資料', 'Query: month（YYYY-MM）', '{ month, events[] }'),
        ('GET', '/api/credit/:userId', '取得信用積分與歷史', 'Path: userId', '{ score, status, history[] }'),
        ('POST', '/api/credit/log', '記錄信用積分變動', '{ user_id, action, change, score_after }', '{ success, logged_at }'),
        ('GET', '/api/campus/buildings', '取得校園建築列表（含無障礙資訊）', 'None', '[ { id, name, accessible, elevator } ]'),
        ('GET', '/api/map/building/:id', '即時查詢建築維修狀態', 'Path: id', '{ id, name, status, maintenance_note }'),
        ('POST', '/api/ai/navigate', 'AI 校園導覽查詢（Dify）', '{ message: string }', '{ answer, confidence, source }'),
        ('POST', '/api/ai/regulations', 'AI 法規 RAG 查詢（Dify）', '{ query: string }', '{ answer, sources[], confidence }'),
        ('POST', '/api/ai/venue-workflow', '場地預約流程 RAG 查詢', '{ requirement: string }', '{ workflow_steps[], recommendation }'),
        ('POST', '/api/ai/screen', 'AI 智慧預審（Dify Workflow）', '{ user_id, role, reason, attached_files }', '{ risk_level, reasoning, suggested_tags, action }'),
        ('POST', '/api/ai/generate-plan', 'AI 活動企劃生成器（Dify）', '{ event_name, event_type, expected_participants }', '{ title, sections[], compliance_check, estimated_budget }'),
        ('POST', '/api/notifications/track', '記錄通知已讀回條（含 IP/時間戳）', '{ token, timestamp }', '{ success, token, ip, tracked_at }'),
        ('GET', '/api/users', '取得用戶列表（管理員）', 'Query: role, page', '[ { id, name, student_id, role, credit_score } ]'),
    ]
    for i, row_data in enumerate(apis):
        add_table_row(api_table, list(row_data), i % 2 == 1)

    doc.add_paragraph()

    # 3.4 System Architecture Diagrams (Mermaid)
    add_heading(doc, '四、系統架構與流程圖（Mermaid UML）', 2)

    # Diagram 1: System Global Architecture
    add_mermaid_block(doc, '一、系統全域架構圖（含 Dify AI 中台 + Cloudflare 防護層）',
    '''graph TB
    subgraph 用戶端 ["用戶端瀏覽器"]
        A["🌐 Vue 3 SPA<br/>GSAP + Leaflet.js<br/>i18n 5國語言"]
    end

    subgraph CF ["☁️ Cloudflare Edge Layer"]
        B["🛡️ WAF<br/>@fju.edu.tw Skip Rules"]
        C["🤖 Turnstile<br/>非互動驗證"]
        D["💾 R2 Storage<br/>圖片/PDF/文件"]
        E["⚡ Workers AI<br/>Llama-3 意圖過濾"]
    end

    subgraph BE ["🖥️ Laravel Backend (PHP 8.3)"]
        F["🔑 Auth Module<br/>Google OAuth + hd 檢查<br/>2FA PHPGangsta + Redis"]
        G["📋 Reservation Module<br/>三階段預約流程<br/>悲觀鎖 + Transaction"]
        H["🏛️ Resource Module<br/>場地/器材管理<br/>LINE/SMS 通知"]
        I["📊 Dashboard Module<br/>五角色儀表板<br/>統計 API"]
        J["👁️ Observer Pattern<br/>信用分<60觸發JWT失效"]
    end

    subgraph AI ["🤖 Dify AI 中台"]
        K["📚 Knowledge Base<br/>RAG 引擎<br/>Pinecone 向量資料庫"]
        L["⚙️ Workflow Engine<br/>企劃生成 + 智慧預審<br/>GuzzleHTTP 呼叫"]
        M["💬 Chat API<br/>AI 導覽/申訴摘要<br/>WebSocket"]
    end

    subgraph DB ["🗄️ 資料層"]
        N[("MySQL 8.0<br/>Materialized View")]
        O[("Redis<br/>2FA + 計時 + Session")]
    end

    subgraph NOTIF ["📨 多通路通知鏈"]
        P["📱 LINE Notify"]
        Q["💬 SMS 簡訊"]
        R["📧 SMTP 收據"]
        S["📮 Outlook Graph API"]
    end

    A -->|HTTPS| B --> C --> BE
    BE -->|GuzzleHTTP| AI
    BE --> DB
    BE --> NOTIF
    CF --> D
    CF --> E
    BE -->|上傳| D
    K --> N
    ''')

    # Diagram 2: Three-phase Reservation State Machine
    add_mermaid_block(doc, '二、三階段資源調度狀態圖（志願序演算法 + 衝突協商）',
    '''stateDiagram-v2
    [*] --> 提交申請: 用戶填寫表單
    提交申請 --> AI_預審: GuzzleHTTP → Dify Workflow

    AI_預審 --> 自動核准: risk_level = Low
    AI_預審 --> 人工審核: risk_level = Medium
    AI_預審 --> 自動拒絕: risk_level = High

    自動核准 --> 衝突檢測: 系統查詢時段衝突
    人工審核 --> 行政審核中: 課指組人員審核
    行政審核中 --> 衝突檢測: 審核通過
    行政審核中 --> 自動拒絕: 審核拒絕

    衝突檢測 --> Phase3_核定: 無衝突
    衝突檢測 --> 協商階段: Level-1優先 或 時段重疊

    state 協商階段 {
        LINE通知 --> 對話室開啟
        對話室開啟 --> 計時開始: Redis 計時器
        計時開始 --> GPT4介入: 3分鐘靜默
        GPT4介入 --> 達成共識: 雙方接受建議
        GPT4介入 --> 強制結束: 6分鐘逾時
        強制結束 --> 扣除信用: -10分 + credit_logs記錄
        扣除信用 --> 重新分配: 系統依權重遞補
    }

    達成共識 --> Phase3_核定: 雙方確認
    Phase3_核定 --> 核定通過: 產出PDF申請單 + TOTP QR Code
    核定通過 --> [*]: 通知所有相關人員
    自動拒絕 --> [*]: 發送拒絕通知
    ''')

    # Diagram 3: Multi-role Swimlane
    add_mermaid_block(doc, '三、多角色全端功能泳道圖（登入→搜尋→預約→審核→簽到）',
    '''sequenceDiagram
    actor S as 👨‍🎓 學生
    actor O as 🎭 社團幹部
    actor P as 👨‍🏫 指導教授
    actor A as ⚙️ 行政人員
    participant SYS as 🖥️ FJU Smart Hub

    Note over S,SYS: 階段一：身份驗證
    S->>SYS: Google OAuth 登入（@cloud.fju.edu.tw）
    SYS-->>S: 2FA TOTP 驗證碼請求
    S->>SYS: 輸入 TOTP 代碼
    SYS-->>S: JWT Cookie 設置，進入主頁

    Note over O,SYS: 階段二：場地申請（社團幹部）
    O->>SYS: 選擇場地 + 填寫活動說明（志願序申請）
    SYS->>SYS: AI 智慧預審（Dify RAG）
    SYS-->>O: risk_level = Medium + 警示標記
    A->>SYS: 查看 AI 預審結果，人工複審
    A-->>SYS: 核准（狀態→APPROVED）
    SYS-->>O: LINE Notify + SMTP 通知核准

    Note over S,SYS: 階段三：簽到 + 信用積分
    S->>SYS: 活動簽到（QR Code 掃描）
    SYS->>SYS: 更新出席紀錄，信用 +2 分
    SYS-->>S: 信用積分更新通知

    Note over P,SYS: 補充：教授查看職能追蹤
    P->>SYS: 查看職能雷達圖 + 活動熱力圖
    SYS-->>P: 學生成長數據視覺化展示
    ''')

    # Diagram 4: Dify AI RAG Flow
    add_mermaid_block(doc, '四、Dify AI 智慧預審與 RAG 流程圖',
    '''flowchart TD
    A["📝 用戶提交預約申請\n{user_id, role, reason, attached_files}"]
    B["🔗 Laravel GuzzleHTTP\nPOST /v1/workflows/run"]
    C["🤖 Dify Workflow Engine"]

    subgraph RAG ["📚 RAG 知識庫檢索"]
        D["Pinecone 向量搜尋\n相似度 top-k 法規條文"]
        E["輔大課指組法規庫\n- 場地設備借用管理辦法\n- 學生社團活動管理辦法\n- 會議紀錄/FAQ"]
    end

    F{"🧠 LLM 風險評估\ngpt-4 / claude-3"}

    G["✅ Low Risk\nAuto Approve"]
    H["⚠️ Medium Risk\nPending Manual Review"]
    I["❌ High Risk\nAuto Reject"]

    J["📊 回傳 JSON 結果\n{task_id, risk_level, reasoning,\nsuggested_tags, action, law_reference}"]

    A --> B --> C
    C --> D --> E
    D --> F
    F -->|合規| G --> J
    F -->|需審查| H --> J
    F -->|違規| I --> J
    J --> K["Laravel 存入 ai_screening_logs\n更新 reservation.status"]
    K --> L["前端展示 AI 預審卡片\n（含 law_reference 引用）"]
    ''')

    # Diagram 5: Security & Identity Lifecycle
    add_mermaid_block(doc, '五、安全驗證與身分生命週期圖',
    '''flowchart LR
    A["🌐 使用者存取系統"]
    B["☁️ Cloudflare WAF\n@fju.edu.tw Skip Rules"]
    C["🤖 Turnstile 驗證\n預約/申復入口"]

    D{"Google OAuth\nhd 域名檢查"}
    E["❌ 非 @fju.edu.tw\n拒絕進入"]
    F["✅ 域名通過"]

    G["2FA 驗證\nTOTP (PHPGangsta)\n或 SMS (Redis 快取 5分鐘)"]
    H["❌ 2FA 失敗\n鎖定帳號 15 分鐘"]
    I["✅ 2FA 通過"]

    J["JWT 產出\nHttpOnly Cookie\nRS256 簽章"]

    subgraph OBSERVER ["🔍 觀察者模式 (Observer Pattern)"]
        K["監控 credit_score"]
        L{"credit_score < 60?"}
        M["INVALIDATE_JWT\n強制登出 + 停權"]
        N["LINE Notify + SMS\n通知用戶停權原因"]
    end

    O["📋 防抵賴追蹤\nUnique Token → track.php\n記錄時間戳 + IP"]

    A --> B --> C
    B --> D
    D -->|hd != fju.edu.tw| E
    D -->|hd = fju.edu.tw| F
    F --> G
    G -->|失敗| H
    G -->|成功| I
    I --> J
    J --> K --> L
    L -->|是| M --> N
    L -->|否| O
    ''')

    doc.add_paragraph()

    # 3.3 Resource Requirements
    add_heading(doc, '五、資源需求', 2)

    add_heading(doc, '開發與營運經費預估表', 3)

    add_paragraph(doc, '▌ 開發費用', bold=True, color=BLUE)
    dev_table = doc.add_table(rows=1, cols=4)
    dev_table.style = 'Table Grid'
    add_table_header(dev_table, ['項目', '說明', '數量/單位', '費用（NT$）'])
    dev_costs = [
        ('硬體設備', '開發用筆電 (MacBook Pro / ThinkPad)', '2台', '120,000'),
        ('雲端伺服器', 'Cloudflare Pages + Workers（免費方案）', '1年', '0'),
        ('MySQL 資料庫', 'PlanetScale/Cloudflare D1（免費方案）', '1年', '0'),
        ('Dify AI 平台', 'Dify Cloud Professional Plan', '1年', '48,000'),
        ('Redis Cloud', 'Redis Upstash 免費方案 + 升級', '1年', '6,000'),
        ('Cloudflare R2', 'R2 存儲 10GB（免費）', '1年', '0'),
        ('LINE Notify API', '免費（需申請 LINE Business）', '1年', '0'),
        ('SMS 服務', '三竹/中華電信 SMS API', '1000則', '3,000'),
        ('人力成本（工讀生）', '系統開發工讀生 2 名，每月 160 小時', '6個月', '96,000'),
        ('指導費', '業師諮詢費', '10小時', '15,000'),
        ('測試費', '系統測試（壓力測試/滲透測試）', '1次', '20,000'),
    ]
    subtotal_dev = 308000
    for i, row_data in enumerate(dev_costs):
        add_table_row(dev_table, list(row_data), i % 2 == 1)
    total_row = dev_table.add_row()
    total_row.cells[0].text = '開發費用合計'
    total_row.cells[3].text = f'{subtotal_dev:,}'
    set_cell_bg(total_row.cells[0], 'DAA520')
    set_cell_bg(total_row.cells[3], 'DAA520')
    for cell in total_row.cells:
        set_cell_bg(cell, 'FDB913')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    doc.add_paragraph()
    add_paragraph(doc, '▌ 三年營運費用預估', bold=True, color=BLUE)
    ops_table = doc.add_table(rows=1, cols=5)
    ops_table.style = 'Table Grid'
    add_table_header(ops_table, ['項目', '第1年（NT$）', '第2年（NT$）', '第3年（NT$）', '說明'])
    ops_costs = [
        ('Cloudflare Pages Pro', '0', '6,000', '6,000', '第2年起升級方案'),
        ('Dify AI 平台', '48,000', '48,000', '48,000', '年費方案'),
        ('域名維護', '500', '500', '500', '.edu.tw 由學校統一管理'),
        ('SSL 憑證', '0', '0', '0', 'Cloudflare 免費提供'),
        ('系統維護', '24,000', '24,000', '24,000', '工讀生每月 2 小時維護'),
        ('SMS 費用', '3,000', '5,000', '7,000', '預估用量成長'),
    ]
    for i, row_data in enumerate(ops_costs):
        add_table_row(ops_table, list(row_data), i % 2 == 1)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # CHAPTER 4: REVIEW
    # ══════════════════════════════════════════════
    add_heading(doc, '第四章　系統專題實作檢討', 1)

    add_heading(doc, '一、發展中遭遇到問題、困難與解決方法', 2)

    problems = [
        {
            'problem': '問題一：場地衝突協商的即時性問題',
            'desc': '初期採用傳統 HTTP Polling 實現協商對話，導致訊息延遲 2-5 秒，用戶體驗不佳。',
            'solution': '改用 Socket.io WebSocket 實現雙向即時通訊，並將 3/6 分鐘計時器移至 Redis 管理，確保多台伺服器間計時一致性。效果：訊息延遲降至 < 100ms。'
        },
        {
            'problem': '問題二：AI 預審準確率不穩定',
            'desc': 'Dify RAG 初期因知識庫文件格式不統一（混合 PDF/Word/網頁），導致法規檢索準確率僅 65%。',
            'solution': '統一採用 JSON 格式建立 Knowledge Base（符合 source_id、category、content、metadata 規範），並對每份法規文件進行人工標注重要條文。準確率提升至 88%。'
        },
        {
            'problem': '問題三：MySQL 並發預約 Race Condition',
            'desc': '多人同時預約同一場地時，系統偶發重複核准問題（同一時段被兩組用戶同時預約成功）。',
            'solution': '改用 MySQL InnoDB `SELECT ... FOR UPDATE` 悲觀鎖機制，並包裹在 `START TRANSACTION / COMMIT` 中確保原子性。透過壓力測試（JMeter 100 並發）驗證解決。'
        },
        {
            'problem': '問題四：Cloudflare Pages 靜態檔案服務問題',
            'desc': '初期直接使用 Node.js `fs` 模組讀取靜態檔案，部署至 Cloudflare Workers 後因無檔案系統支援而報錯。',
            'solution': '改用 `hono/cloudflare-workers` 的 `serveStatic` 方法，將所有靜態資源放置於 `public/` 目錄，透過 Cloudflare Pages 的靜態資源服務機制處理。'
        },
        {
            'problem': '問題五：五國語言 i18n 切換不完整',
            'desc': '初版 i18n 僅切換部分 UI 元素，動態渲染的 JavaScript 組件（如地圖 Popup、Modal 內容）未能即時更新語言。',
            'solution': '建立全域 `refreshUILabels()` 函式搭配 `data-i18n` 屬性選擇器，確保所有 UI 元素（包含動態生成內容）在語言切換時完整更新。',
        },
    ]
    for item in problems:
        p = doc.add_paragraph()
        run = p.add_run(item['problem'])
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE
        p.paragraph_format.space_after = Pt(2)

        add_paragraph(doc, f'問題描述：{item["desc"]}', size=9, color=RED, indent=0.5)
        add_paragraph(doc, f'解決方案：{item["solution"]}', size=9, color=GREEN, indent=0.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 4.2 SWOT
    add_heading(doc, '二、系統優缺點（SWOT）評估', 2)
    swot_table = doc.add_table(rows=2, cols=2)
    swot_table.style = 'Table Grid'

    swot_data = [
        ('優勢（Strengths）',
         '・整合 AI（Dify RAG）實現智慧化行政流程\n・獨特的 3/6 分鐘衝突協商機制\n・五國語言支援覆蓋廣泛用戶群\n・Cloudflare 邊緣部署確保低延遲\n・完整的信用積分系統確保公平性\n・E-Portfolio 功能具高附加值',
         '劣勢（Weaknesses）',
         '・依賴外部 Dify API（服務中斷風險）\n・地圖資料需人工定期維護更新\n・RAG 準確率仍有 12% 誤差空間\n・Cloudflare Workers CPU 時間限制（10ms）\n・手機版介面尚未完全優化'),
        ('機會（Opportunities）',
         '・可推廣至全台大學（SaaS 模式）\n・AI 技術快速發展，模型能力持續提升\n・大學數位轉型政策支持\n・可整合教育部 SDGs 評估需求\n・未來可拓展至校友網絡管理',
         '威脅（Threats）',
         '・個資保護法規對 AI 數據使用的限制\n・競爭者可能推出類似功能\n・Dify/Cloudflare 服務費用變動風險\n・學生對 AI 決策信任度不足\n・校方 IT 安全政策可能限制部分功能'),
    ]

    for row_idx in range(2):
        for col_idx in range(2):
            cell = swot_table.cell(row_idx, col_idx)
            title = swot_data[row_idx][col_idx * 2]
            content = swot_data[row_idx][col_idx * 2 + 1]

            # Set cell background
            bg_colors = [['D4EDDA', 'FFF3CD'], ['D1ECF1', 'F8D7DA']]
            set_cell_bg(cell, bg_colors[row_idx][col_idx])

            p = cell.paragraphs[0]
            run = p.add_run(title + '\n')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = BLUE
            run2 = p.add_run(content)
            run2.font.size = Pt(8.5)
            run2.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    # 4.3 Development Insights
    add_heading(doc, '三、發展心得', 2)
    add_paragraph(doc, '''
本次專題開發 FJU Smart Hub，是一次從零到一整合 AI、雲端服務與傳統行政流程的完整實戰經驗。

透過這次開發，我深刻體會到「技術選型」的重要性。一開始選擇了過於複雜的技術堆疊（試圖在單一 Cloudflare Worker 中運行完整 Laravel 框架），後來意識到應以「邊緣優先（Edge-First）」原則重新設計架構，將 API 邏輯分層處理，前端采用 Hono 輕量框架服務靜態資源，後端 Laravel 負責複雜業務邏輯，大幅提升了系統的可維護性。

AI 整合部分，Dify RAG 的導入讓我理解到 AI 並非萬能 — 知識庫的品質比模型的能力更加關鍵。花費大量時間整理、標注輔大法規文件，最終使預審準確率從 65% 提升至 88%，這個過程讓我理解到「數據工程」在 AI 系統中的核心地位。

衝突協商機制的設計是本系統最有創意的部分。透過 Redis 精確計時、WebSocket 即時通訊與 GPT-4 的結構化建議，我們將原本需要 1-3 天電話協調的場地衝突問題，縮短至最多 6 分鐘解決，這讓我對「技術解決社會問題」有了更具體的理解。
    '''.strip())

    # 4.4 Future Outlook
    add_heading(doc, '四、未來展望', 2)
    future_items = [
        '移動端 App（React Native）：將現有 PWA 升級為原生 App，提供推播通知與離線完整支援',
        '進階 AI 功能：整合多模態 AI（圖片識別），讓 AI 可分析活動現場照片自動填寫成果報告',
        '區塊鏈學習履歷：將 E-Portfolio 關鍵成就上鏈（NFT 幹部證書），確保履歷不可竄改',
        'SDGs 自動評估：透過 AI 分析活動內容，自動分類對應聯合國永續發展目標（SDGs 17項）',
        '跨校園聯盟：建立輔大與鄰近大學（輔仁、臺北大學等）的社團交流平台，擴大系統影響力',
        '物聯網整合：在場地入口安裝 NFC/RFID 讀取器，實現無縫自動簽到功能',
        'Mapbox 升級：將 Leaflet.js 地圖升級為 Mapbox GL JS，支援 3D 建築渲染與空間競爭熱圖',
    ]
    for item in future_items:
        add_bullet(doc, item)

    doc.add_page_break()

    # ══════════════════════════════════════════════
    # APPENDIX
    # ══════════════════════════════════════════════
    add_heading(doc, '附錄', 1)

    # Sprint Progress
    add_heading(doc, '一、Sprint 進度報告', 2)

    sprints = [
        {
            'num': 1,
            'goal': '建立系統基礎架構與認證機制',
            'duration': '2026/01/10 – 2026/01/24（2週）',
            'stories': ['US-13（使用者身份驗證）', 'US-16（資料庫 Schema 設計）'],
            'done': ['Google OAuth + hd 域名驗證', '2FA TOTP 整合（PHPGangsta）', 'JWT HttpOnly Cookie 實作', 'MySQL Schema v1.0（15個資料表）', 'Cloudflare Pages + Hono 部署'],
            'issue': '2FA SMS 服務商 API 整合費時，改用 Redis Mock 完成測試',
            'review': '基礎架構完整，所有核心安全機制已就位',
            'retro': '下個 Sprint 需要加快 AI 模組開發速度',
        },
        {
            'num': 2,
            'goal': '實作場地預約三階段流程與 AI 預審',
            'duration': '2026/01/25 – 2026/02/14（3週）',
            'stories': ['US-01（場地預約）', 'US-03（AI 預審）'],
            'done': ['三階段預約 UI（志願序→協商→核定）', 'Dify Workflow API 整合', 'AI 預審 risk_level 三段式判斷', 'MySQL 悲觀鎖防止 Race Condition', 'PDF 申請單自動產出'],
            'issue': 'Dify API 速率限制（Rate Limit）導致測試緩慢，實作本地 Mock 繞過',
            'review': 'AI 預審核心功能完成，準確率 65%（待優化）',
            'retro': 'Knowledge Base 品質需要大幅改善',
        },
        {
            'num': 3,
            'goal': '完成衝突協商機制與信用系統',
            'duration': '2026/02/15 – 2026/03/01（2週）',
            'stories': ['US-04（衝突協商）', 'US-11（信用系統）'],
            'done': ['WebSocket 即時對話室（Socket.io）', 'Redis 3/6 分鐘精確計時', 'GPT-4 介入機制（3 個建議）', '紅光扣分動畫', '信用觀察者模式（<60分強制登出）'],
            'issue': 'WebSocket 在 Cloudflare Pages 環境不支援，改用 Long Polling 前端模擬',
            'review': '協商機制完整實作，用戶測試反應正面',
            'retro': '需要處理更多邊緣案例（如一方斷線）',
        },
        {
            'num': 4,
            'goal': '完成所有 UI 介面與 10 大功能模組',
            'duration': '2026/03/02 – 2026/03/20（3週）',
            'stories': ['US-02（互動地圖）', 'US-07（E-Portfolio）', 'US-09（AI 企劃生成）'],
            'done': ['Leaflet.js 互動地圖 + 無障礙標記', 'GSAP 玻璃行事曆面板', 'E-Portfolio 職能雷達圖', 'AI 企劃生成器（Dify Workflow）', '幹部證書自動化', '動態活動牆', '5 國語言 i18n'],
            'issue': '地圖效能問題（同時渲染 10+ 標記時卡頓），改用 Lazy Loading',
            'review': '全部 10 大功能模組完成，系統完整度達 95%',
            'retro': 'Landing Page 動畫效果還需加強',
        },
        {
            'num': 5,
            'goal': '優化 Landing Page 動畫與撰寫 SA 規格書',
            'duration': '2026/03/21 – 2026/04/03（2週）',
            'stories': ['US-15（Landing Page 動畫）', 'US-17（SA 文件）'],
            'done': ['GSAP ScrollTrigger 滾動動畫', '活動輪播 Carousel', '痛點三欄式卡片', 'Bento Grid 功能展示', '圓形頭像評價輪播', '即時快訊橫幅 Ticker', 'SA 規格書 v5.0 完成'],
            'issue': '無',
            'review': '系統完整交付，Landing Page 視覺效果優秀',
            'retro': '專案成功，未來可推廣至其他大學',
        },
    ]

    for sprint in sprints:
        p = doc.add_paragraph()
        run = p.add_run(f'Sprint {sprint["num"]}：{sprint["goal"]}')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE
        p.paragraph_format.space_after = Pt(2)

        sprint_table = doc.add_table(rows=6, cols=2)
        sprint_table.style = 'Table Grid'
        sprint_labels = ['Sprint 目標期間', '本 Sprint 使用者故事', '完成項目', '遭遇問題', 'Sprint Review', 'Sprint Retrospective']
        sprint_vals = [
            sprint['duration'],
            '・'.join(sprint['stories']),
            '\n・'.join(['・' + d for d in sprint['done']]),
            sprint['issue'],
            sprint['review'],
            sprint['retro'],
        ]
        for i, (label, val) in enumerate(zip(sprint_labels, sprint_vals)):
            row = sprint_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = val
            set_cell_bg(row.cells[0], '003153')
            set_cell_bg(row.cells[1], 'F8FAFC' if i % 2 == 0 else 'FFFFFF')
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = WHITE
                    run.font.size = Pt(8.5)
            for para in row.cells[1].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = DARK_GRAY
        doc.add_paragraph()

    # Contribution Table
    add_heading(doc, '二、分工與貢獻度說明表', 2)
    contrib_table = doc.add_table(rows=1, cols=5)
    contrib_table.style = 'Table Grid'
    add_table_header(contrib_table, ['工作項目', '具體任務', '負責人', '工時（小時）', '貢獻比例'])
    contributions = [
        ('系統架構設計', 'Cloudflare Pages + Hono 架構規劃、DB Schema 設計', 'KY0126', '40', '15%'),
        ('後端 API', 'Laravel PHP 8.3 API、JWT Auth、2FA、Dify 整合', 'KY0126', '80', '30%'),
        ('前端開發', 'Vue 3 + GSAP + Leaflet.js、Landing Page、Dashboard', 'KY0126', '70', '26%'),
        ('AI 模組', 'Dify RAG Knowledge Base、企劃生成器、智慧預審', 'KY0126', '35', '13%'),
        ('資料庫', 'MySQL 8.0 Schema、Materialized View、Seed Data', 'KY0126', '20', '7%'),
        ('測試與部署', '壓力測試、GitHub CI/CD、Cloudflare 部署', 'KY0126', '15', '6%'),
        ('文件撰寫', 'SA 規格書 v5.0（本文件）、README.md', 'KY0126', '10', '4%'),
    ]
    for i, row_data in enumerate(contributions):
        add_table_row(contrib_table, list(row_data), i % 2 == 1)
    total_contrib = contrib_table.add_row()
    total_contrib.cells[0].text = '合計'
    total_contrib.cells[3].text = '270'
    total_contrib.cells[4].text = '101%'  # Note: rounding
    set_cell_bg(total_contrib.cells[0], 'DAA520')
    for cell in total_contrib.cells:
        set_cell_bg(cell, 'FDB913')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    doc.add_paragraph()

    # References
    add_heading(doc, '三、參考資料', 2)
    references = [
        ('輔仁大學課外活動指導組 官方網站', 'https://activity.fju.edu.tw/'),
        ('輔仁大學學生社團場地設備借用管理辦法', 'https://activity.fju.edu.tw/generalServices.jsp?labelID=30'),
        ('Dify AI 官方文件', 'https://docs.dify.ai/'),
        ('Hono Web Framework', 'https://hono.dev/'),
        ('Cloudflare Pages 文件', 'https://developers.cloudflare.com/pages/'),
        ('Leaflet.js 互動地圖', 'https://leafletjs.com/'),
        ('GSAP 動畫庫', 'https://greensock.com/gsap/'),
        ('Pinecone 向量資料庫', 'https://www.pinecone.io/'),
        ('PHPGangsta TOTP Library', 'https://github.com/PHPGangsta/GoogleAuthenticator'),
        ('Cloudflare Turnstile', 'https://www.cloudflare.com/products/turnstile/'),
        ('Socket.io WebSocket', 'https://socket.io/'),
        ('Chart.js 圖表庫', 'https://www.chartjs.org/'),
        ('Vue 3 + vue-i18n', 'https://vue-i18n.intlify.dev/'),
    ]
    for i, (title, url) in enumerate(references):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run1 = p.add_run(f'[{i+1}] ')
        run1.font.bold = True
        run1.font.size = Pt(9)
        run2 = p.add_run(f'{title}. ')
        run2.font.size = Pt(9)
        run3 = p.add_run(url)
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0, 102, 204)
        run3.font.underline = True
        p.paragraph_format.space_after = Pt(4)

    # Save document
    output_path = '/home/user/webapp/FJU_Smart_Hub_SA文件_v5.0.docx'
    doc.save(output_path)
    print(f'✅ SA 文件已成功生成：{output_path}')
    return output_path

if __name__ == '__main__':
    path = generate_sa_document()
    print(f'📄 文件大小：{__import__("os").path.getsize(path):,} bytes')
