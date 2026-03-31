"""
FJU Smart Hub - System Analysis Document Generator
Generates Word SA document following the provided template format.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────
# Helper utilities
# ─────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get(side, 'single'))
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), kwargs.get('color', '003087'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x8f)
    return h

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def bullet(doc, text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(11)
    return p

def note_box(doc, text, label="備註"):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    set_cell_bg(cell, 'E8F0FE')
    p = cell.paragraphs[0]
    run = p.add_run(f"💡 {label}：{text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)
    doc.add_paragraph()

def section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────────────────
# Main Document Builder
# ─────────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default styles
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(11)

    # ─────────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('FJU Smart Hub')
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run('全方位校園管理與資源調度系統')
    subtitle_run.font.size = Pt(22)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x8f)

    doc.add_paragraph()

    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    type_run = type_p.add_run('系統分析與設計規格書 (System Analysis Document)')
    type_run.font.size = Pt(16)
    type_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()
    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        '輔仁大學 資訊管理學系\n'
        '第三十四屆 專題文件\n\n'
        '開發團隊：FJU Smart Hub Dev Team\n'
        '指導教授：陳大中 教授\n'
        '系統版本：v3.0\n'
        '更新日期：2026/03/31'
    )
    info_run.font.size = Pt(13)

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # REVISION HISTORY
    # ─────────────────────────────────────────────
    heading(doc, '修改紀錄', 1)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    headers = ['版本', '日期', '修改說明', '修改人']
    for i, h_text in enumerate(headers):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    rows = [
        ('v1.0', '2026/01/15', '初始文件建立，系統架構規劃', 'Dev Team'),
        ('v2.0', '2026/02/20', '新增 Dify AI 整合、三階段預約流程說明', 'Dev Team'),
        ('v3.0', '2026/03/31', '完整十大功能、五角色儀表板、資料庫 Schema', 'Dev Team'),
    ]
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ─────────────────────────────────────────────
    # TABLE OF CONTENTS (manual)
    # ─────────────────────────────────────────────
    heading(doc, '目錄', 1)
    toc_items = [
        ('第一章 系統描述', '3'),
        ('  一、需求分析與市場探索', '3'),
        ('  二、系統發展目的', '4'),
        ('  三、系統範圍', '5'),
        ('  四、背景知識', '5'),
        ('  五、系統限制', '6'),
        ('第二章 軟體需求規格', '7'),
        ('  一、功能需求（使用者故事）', '7'),
        ('  二、非功能性需求', '10'),
        ('第三章 軟體設計規格', '11'),
        ('  一、系統架構設計', '11'),
        ('  二、資料庫設計', '13'),
        ('  三、介面設計', '17'),
        ('  四、Dify AI 整合設計', '20'),
        ('  五、資源需求（預算經費）', '22'),
        ('第四章 系統專題實作檢討', '24'),
        ('  一、發展中遭遇到問題、困難與解決方法', '24'),
        ('  二、系統優缺點（SWOT）評估', '25'),
        ('  三、發展心得', '26'),
        ('  四、未來展望', '26'),
        ('附　錄', '27'),
        ('  進度報告（Sprint 1-3）', '27'),
        ('  二、參考資料', '30'),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}')
        run.font.size = Pt(11)
        # dots + page number
        tab_stop = p.paragraph_format
        p.add_run(f' {"." * (60 - len(item))} {page}').font.size = Pt(10)

    doc.add_page_break()

    # ═════════════════════════════════════════════
    # CHAPTER 1: SYSTEM DESCRIPTION
    # ═════════════════════════════════════════════
    heading(doc, '第一章 系統描述', 1)

    # Section 1
    heading(doc, '一、需求分析與市場探索', 2)

    heading(doc, '問題陳述', 3)
    body(doc,
        '輔仁大學校園內各類資源（場地、器材）的管理與調度缺乏數位化整合平台。'
        '課外活動指導組每年需處理數百件場地預約申請，仰賴紙本流程，容易發生：',
        indent=True)
    bullet(doc, '時段衝突：同等級社團爭搶熱門場地（焯炤館、進修部演講廳），無法自動化協調')
    bullet(doc, '人工瓶頸：行政人員需逐一審核申請，缺乏 AI 輔助預審機制')
    bullet(doc, '信用缺乏追蹤：社團爽約、器材逾期歸還無法即時扣分並停權')
    bullet(doc, '資料分散：社團成果、幹部職能記錄分散於各處，無法形成結構化 E-Portfolio')
    bullet(doc, '通知滯後：場地核定、器材到期等通知依賴人工電話，缺乏 LINE/SMS 自動推播')

    heading(doc, '利害關係人分析', 3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['角色', '說明', '核心需求']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '1a3a8f')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    stakeholders = [
        ('學生',      '全體在學學生，希望快速預約場地、查詢活動、建立個人職能檔案',
                      '依順序預約場地；建立 E-Portfolio；查看活動牆'),
        ('社團幹部',  '各社團負責人，需管理社員、規劃活動、控管經費',
                      '智慧企劃生成；三階段預約流；幹部傳承機制'),
        ('指導教授',  '監督社團運作，提供學術指導',
                      '雷達圖績效考核；學生職能成長監控；風險預警'),
        ('課指組職員','執行行政審核、場地器材管理、法規執行',
                      'AI 輔助預審；SDGs 貢獻統計；行政時效監控'),
        ('資訊中心',  '維護系統穩定性、資安防護',
                      'API 監控儀表板；WAF 日誌；R2 使用率'),
    ]
    for s in stakeholders:
        row = t.add_row()
        for i, val in enumerate(s):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    heading(doc, '需求蒐集方法', 3)
    body(doc,
        '開發團隊採用以下多元方法收集需求：',
        indent=True)
    bullet(doc, '訪談：深度訪談課外活動指導組職員 3 人，了解行政痛點及現行作業流程')
    bullet(doc, '問卷：向 50 位社團幹部發放線上問卷，收回 43 份有效問卷（回收率 86%）')
    bullet(doc, '觀察：跟訪場地器材借用作業流程一週，記錄手動作業時間與錯誤率')
    bullet(doc, '文件分析：研讀輔仁大學學生社團輔導辦法、場地管理辦法等法規')
    bullet(doc, 'Focus Group：與 8 位不同角色使用者進行焦點團體討論，優先排序功能')

    heading(doc, '競爭者分析', 3)
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['系統名稱', '主要功能', '缺點', '本系統優勢']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '1a3a8f')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    competitors = [
        ('課指組現行系統', '基本場地申請表單', '純紙本，無自動化', '全數位化 + AI 預審'),
        ('HiTeach', '師生互動教學', '僅限課堂互動，無場地預約', '涵蓋社團全生命週期'),
        ('Slido', '線上投票問答', '僅限活動互動功能', '完整行政後台'),
        ('iCal / Google Calendar', '個人行事曆', '無角色權限，無審核流程', '多角色衝突協商機制'),
    ]
    for c in competitors:
        row = t.add_row()
        for i, val in enumerate(c):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    heading(doc, '市場定位 / 利基', 3)
    body(doc,
        'FJU Smart Hub 定位為「輔大校園專屬的智慧資源調度平台」，以下為核心利基：',
        indent=True)
    bullet(doc, '深度整合：唯一結合 Dify AI 智慧預審、志願序演算法、信用懲處狀態機的校園系統')
    bullet(doc, '法規合規：RAG 知識庫直接引用課指組法規，AI 回應帶有精確法條引用')
    bullet(doc, '多角色設計：五種角色專屬儀表板，涵蓋學生、幹部、教授、行政、資訊')
    bullet(doc, '邊緣運算：Cloudflare Pages 全球 CDN 部署，低延遲高可用性')
    bullet(doc, '資安防護：Cloudflare WAF + Turnstile + Google OAuth hd 限制，三層防護')

    section_divider(doc)

    # Section 2
    heading(doc, '二、系統發展目的', 2)

    body(doc, '本系統旨在解決輔仁大學校園資源管理的核心痛點，分角色說明如下：')

    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['角色', '系統對其帶來的價值']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    goals = [
        ('學生',
         '1. 線上預約場地，即時查看可用狀態\n'
         '2. 動態活動牆，依標籤與關鍵字搜尋感興趣活動\n'
         '3. 自動累積 E-Portfolio，匯出 PDF 職能報告\n'
         '4. 即時查看個人信用點數與扣分記錄'),
        ('社團幹部',
         '1. 三階段預約流（申請→協商→核定），減少人工溝通\n'
         '2. AI 企劃書生成器，一鍵產出符合法規的企劃書\n'
         '3. 幹部職能 E-Portfolio 與傳承數位時光膠囊\n'
         '4. 社員成長趨勢、活動滿意度即時分析'),
        ('指導教授',
         '1. 社團績效雷達圖，一目了然掌握社團健康狀況\n'
         '2. 學生職能成長 Spider Chart，觀察能力發展軌跡\n'
         '3. 風險預警系統，信用分 < 60 即時紅燈警示'),
        ('課指組職員',
         '1. AI 智慧預審大幅降低人工審核負擔\n'
         '2. SDGs 貢獻度雷達圖，輔助政策決策\n'
         '3. 行政審核時效漏斗圖，找出流程瓶頸\n'
         '4. 社團參與率趨勢，提供學生事務政策依據'),
        ('資訊中心',
         '1. API 請求成功率與延遲監控\n'
         '2. Cloudflare WAF 攔截記錄視覺化\n'
         '3. R2 存儲空間利用率監控\n'
         '4. 系統負載熱力圖（高峰預測）'),
    ]
    for g in goals:
        row = t.add_row()
        row.cells[0].text = g[0]
        row.cells[1].text = g[1]
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # Section 3
    heading(doc, '三、系統範圍', 2)
    body(doc, '本系統的服務範圍涵蓋以下功能模組：')
    modules = [
        ('場地資源調度', '三階段預約流（志願序申請→衝突協商→官方核定）'),
        ('器材借用追蹤', 'TOTP QR Code 雙重驗證、LINE/SMS 到期提醒'),
        ('社團生命週期管理', '社員管理、活動規劃、幹部傳承、數位時光膠囊'),
        ('動態活動牆', '全校活動公告，支援關鍵字與標籤搜尋'),
        ('E-Portfolio', '職能標籤累積、PDF 匯出、志工時數記錄'),
        ('AI 智慧預審', 'Dify RAG 比對法規、AI 企劃書生成、申訴摘要'),
        ('信用懲處機制', '觀察者模式即時扣分、JWT 強制失效、停權管理'),
        ('多角色儀表板', '五種角色（學生/幹部/教授/行政/資訊）專屬統計圖表'),
        ('全域行事曆', '多社團活動整合、ICS 匯出'),
        ('2FA 雙因素驗證', '仲裁/核銷/後台登入強制 TOTP 驗證'),
    ]
    for m, desc in modules:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{m}：')
        run.font.bold = True
        run.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    body(doc, '系統範圍外（不在本版本開發計畫內）：')
    bullet(doc, '宿舍管理系統整合（另案處理）')
    bullet(doc, '選課系統直接串接（暫以間接方式實現）')
    bullet(doc, '跨校際聯合活動平台（未來版本擴充）')

    # Section 4
    heading(doc, '四、背景知識', 2)

    heading(doc, 'Cloudflare Pages / Workers', 3)
    body(doc,
        'Cloudflare Pages 是全球邊緣運算平台，支援在距離用戶最近的節點執行程式碼。'
        '本系統後端採用 Hono 框架部署於 Cloudflare Workers，享有免費方案下每日 100,000 次請求額度、'
        'D1 SQLite 資料庫（1GB）、KV 儲存（1GB）等服務，適合校園規模系統。',
        indent=True)

    heading(doc, 'Dify AI 中台', 3)
    body(doc,
        'Dify 是開源的 LLM 應用開發平台，提供 RAG（檢索增強生成）知識庫、Workflow 工作流、'
        'API 串接等能力。本系統將課指組法規文件建立為 Dify Knowledge Base，'
        '透過 GuzzleHTTP / Fetch API 呼叫 Dify Workflow，實現 AI 智慧預審及企劃書生成。',
        indent=True)

    heading(doc, 'Vue 3 前端框架', 3)
    body(doc,
        'Vue 3 採用 Composition API，支援響應式資料綁定與組件化開發。'
        '搭配 vue-i18n 實現中/英/日/韓四國語系切換，ECharts / Chart.js 繪製多種統計圖表，'
        'GSAP ScrollTrigger 實現流暢入場動畫。',
        indent=True)

    heading(doc, 'PHP Laravel 後端', 3)
    body(doc,
        'Laravel 12（PHP 8.3+）採用 MVC 架構，透過 Eloquent ORM 操作 MySQL 8.0（InnoDB）。'
        '資料庫採用悲觀鎖（SELECT FOR UPDATE）防止高頻預約 Race Condition，'
        'Materialized View 快取統計查詢，提升儀表板回應速度。',
        indent=True)

    # Section 5
    heading(doc, '五、系統限制', 2)
    bullet(doc, '身份限制：僅允許 @cloud.fju.edu.tw 與 @mail.fju.edu.tw 網域帳號登入，由 Google OAuth hd 參數強制檢查')
    bullet(doc, '效能限制：Cloudflare Free Plan 每個請求 CPU 時間上限 10ms，AI 請求採異步 outbox_table 機制')
    bullet(doc, '儲存限制：D1 資料庫 1GB 上限；R2 物件儲存 10GB 免費額度，圖片須經 Canvas 壓縮')
    bullet(doc, '即時通訊限制：WebSocket 長連接功能不在 Cloudflare Pages 部署範圍，採 Server-Sent Events 替代')
    bullet(doc, '法規限制：場地預約最終仍須按課指組規定完成紙本留存，系統負責「網路審核先行」')
    bullet(doc, '語言限制：目前 AI 生成結果以正體中文為主，英文模式精準度較低（待優化）')

    note_box(doc,
        '未來版本預計整合 Cloudflare AI Workers（Llama-3）進行客戶端快速語意過濾，降低對外部 Dify API 的依賴。')

    doc.add_page_break()

    # ═════════════════════════════════════════════
    # CHAPTER 2: SOFTWARE REQUIREMENTS
    # ═════════════════════════════════════════════
    heading(doc, '第二章 軟體需求規格', 1)

    heading(doc, '一、功能需求', 2)

    body(doc, '使用者角色說明')
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['角色', '說明']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    roles = [
        ('學生（Student）', '全校在學學生；可預約場地/器材、瀏覽活動牆、建立 E-Portfolio、查看信用點數'),
        ('社團幹部（Club Officer）', '各社團負責人；除學生權限外，可管理社員、申請大型活動、使用 AI 企劃生成器'),
        ('指導教授（Professor）', '社團顧問；可查看所指導社團績效、監控學生職能成長、接收風險預警'),
        ('課指組職員（Admin）', '課外活動指導組行政人員；可審核預約申請、管理法規 FAQ、查看全局統計'),
        ('資訊中心（IT Admin）', '系統管理員；可查看系統監控、WAF 日誌、管理用戶帳號、設定系統參數'),
    ]
    for r in roles:
        row = t.add_row()
        row.cells[0].text = r[0]
        row.cells[1].text = r[1]
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    body(doc, '使用者故事對應')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('優先順序說明：1 = Sprint 1、2 = Sprint 2、3 = Sprint 3').font.size = Pt(10)

    doc.add_paragraph()
    body(doc, '使用者故事卡', indent=False)

    # User Story Tables
    user_stories = [
        {
            'id': '1', 'type': 'Epic', 'title': '場地資源調度',
            'role': '社團幹部', 'want': '透過系統進行三階段場地預約（申請→協商→核定）',
            'value': '減少人工溝通，提升預約效率，降低時段衝突',
            'children': [
                ('1.1', '填寫預約申請', '社團幹部', '填寫場地預約表單並設定志願序', '提交申請給 AI 預審', '1'),
                ('1.2', 'AI 智慧預審', '系統', '自動透過 Dify RAG 比對法規並回傳風險等級', '提前識別違規申請，減少人工負擔', '1'),
                ('1.3', '衝突協商通知', '社團幹部', '當偵測到時段衝突時收到 LINE Notify 協商邀請', '透過平台完成自主協調，不需等待行政介入', '2'),
                ('1.4', '核定確認與 TOTP', '課指組職員', '核准後系統自動生成 30 秒即逝 TOTP QR Code', '雙重驗證，防止代管糾紛', '2'),
                ('1.5', '信用連動', '系統', '未依約簽到時自動扣除信用點數，低於 60 分觸發停權', '維護系統公平性，鼓勵負責任使用', '1'),
            ]
        },
        {
            'id': '2', 'type': 'Epic', 'title': '器材借用與追蹤',
            'role': '學生', 'want': '線上申請借用校內器材並查看歸還期限',
            'value': '方便快速取得所需器材，避免逾期遺忘歸還',
            'children': [
                ('2.1', '查詢器材庫存', '學生', '查看各器材目前可借數量與位置', '規劃借用計畫', '1'),
                ('2.2', '申請借用', '學生/社團幹部', '線上填寫借用申請並選擇歸還日期', '省去人工登記時間', '1'),
                ('2.3', 'LINE/SMS 到期提醒', '系統', '自動於到期前 3 天發送歸還提醒通知', '降低逾期率', '2'),
                ('2.4', '盤點報表', '課指組職員', '查看各器材出借率、損耗率統計', '輔助採購與維修決策', '3'),
            ]
        },
        {
            'id': '3', 'type': 'Epic', 'title': 'E-Portfolio 職能檔案',
            'role': '學生', 'want': '系統自動累積參與活動、志工時數、獲得證書的記錄',
            'value': '建立完整的校園職能履歷，畢業時匯出 PDF',
            'children': [
                ('3.1', '自動記錄活動參與', '系統', '每次簽到後自動新增活動參與記錄', '減少手動輸入', '1'),
                ('3.2', '技能標籤管理', '學生', '自行新增或由系統 AI 建議職能標籤', '精準描述個人專長', '2'),
                ('3.3', '幹部證書自動化', '社團幹部', '解散時系統自動生成帶數位簽章的幹部證書', '取代人工製作', '2'),
                ('3.4', 'PDF 匯出', '學生', '一鍵產出完整 E-Portfolio PDF 報告', '求職/申請研究所時可附上', '2'),
            ]
        },
        {
            'id': '4', 'type': 'Epic', 'title': 'Dify AI 工具中心',
            'role': '社團幹部/課指組職員', 'want': '使用 AI 工具輔助日常行政工作',
            'value': '大幅降低行政負擔，提升決策品質',
            'children': [
                ('4.1', 'AI 企劃書生成', '社團幹部', '輸入活動基本資訊，AI 自動生成符合法規的企劃書草稿', '省時省力，法規合規', '1'),
                ('4.2', 'AI 申訴摘要', '課指組職員', '上傳申訴文件，AI 提取重點並分析風險', '提升審核效率', '2'),
                ('4.3', 'RAG 法規問答', '全體用戶', '自然語言查詢相關法規，AI 引用具體條文回答', '降低法規違規風險', '2'),
            ]
        },
        {
            'id': '5', 'type': 'Epic', 'title': '信用點數與安全機制',
            'role': '系統', 'want': '實作完整的信用點數追蹤與 JWT 安全機制',
            'value': '確保系統公平性，防止惡意使用',
            'children': [
                ('5.1', '信用點數扣除', '系統', '爽約/逾期/違規自動扣分，管理員可手動調整', '維護系統公平', '1'),
                ('5.2', 'JWT 即時失效', '系統', '信用分 < 60 時觀察者模式觸發 JWT 版本更新，強制重新登入', '即時停權，防止繼續使用', '1'),
                ('5.3', '2FA 強制驗證', '管理員/幹部', '仲裁/核銷/後台登入強制 TOTP 二次驗證', '防止帳號被盜用', '1'),
                ('5.4', 'Google OAuth hd 檢查', '系統', '登入時強制驗證 hd == cloud.fju.edu.tw', '確保只有輔大帳號可登入', '1'),
            ]
        },
    ]

    for epic in user_stories:
        # Epic card
        t = doc.add_table(rows=2, cols=4)
        t.style = 'Table Grid'
        header_cell = t.cell(0, 0)
        header_cell.merge(t.cell(0, 3))
        set_cell_bg(header_cell, 'C8A951')
        run = header_cell.paragraphs[0].add_run(f"{epic['id']} {epic['title']}")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 48, 135)

        r1c0 = t.cell(1, 0)
        r1c0.text = '名稱'
        r1c1 = t.cell(1, 1)
        r1c1.merge(t.cell(1, 3))
        r1c1.text = f"{epic['id']} {epic['title']}"
        for cell in [r1c0, r1c1]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

        # Epic story
        t2 = doc.add_table(rows=3, cols=4)
        t2.style = 'Table Grid'
        labels = ['角色', '需求', '價值', '類別']
        vals = [
            f"身為 {epic['role']}",
            f"我想要 {epic['want']}",
            f"因此我可以 {epic['value']}",
            'Epic'
        ]
        for i, (lbl, val) in enumerate(zip(labels, vals)):
            t2.cell(0, i).text = lbl
            t2.cell(1, i).text = val
            t2.cell(2, i).text = ''
            for row in range(3):
                for para in t2.cell(row, i).paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

        # User Story cards
        for story in epic['children']:
            sid, sname, srole, swant, svalue, sprint = story
            t3 = doc.add_table(rows=1, cols=4)
            t3.style = 'Table Grid'
            h_cell = t3.cell(0, 0)
            h_cell.merge(t3.cell(0, 3))
            set_cell_bg(h_cell, 'E8F0FE')
            run = h_cell.paragraphs[0].add_run(f"  {sid} {sname}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 48, 135)

            t4 = doc.add_table(rows=3, cols=4)
            t4.style = 'Table Grid'
            story_labels = ['名稱', sid + ' ' + sname, '類別', 'User Story']
            for i, (col0, col1) in enumerate(zip(['名稱','角色','需求','價值'], [sid+' '+sname, f'身為 {srole}', f'我想要 {swant}', f'因此我可以 {svalue}'])):
                t4.cell(0 if i < 2 else 1, 0 if i % 2 == 0 else 2).text = col0
                t4.cell(0 if i < 2 else 1, 1 if i % 2 == 0 else 3).text = col1

            priority_row = t4.add_row()
            priority_row.cells[0].text = '優先順序'
            priority_row.cells[1].text = sprint
            for row in t4.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(9)

            doc.add_paragraph()

    heading(doc, '二、非功能性需求', 2)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['類別', '需求描述', '驗收標準']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    nfr = [
        ('效能', 'API 回應時間 ≤ 200ms（p95）', 'Cloudflare Analytics 監控，p95 < 200ms'),
        ('可用性', '系統可用率 ≥ 99.5%', 'Cloudflare Pages SLA 99.9%，每月允許停機 < 3.6hr'),
        ('安全性', 'SQL Injection / XSS 防護', 'OWASP ZAP 掃描 0 高危漏洞'),
        ('安全性', '僅限 @cloud.fju.edu.tw 登入', 'Google OAuth hd 參數強制驗證'),
        ('安全性', '重要操作 2FA 驗證', 'PHPGangsta + Redis TOTP，30秒有效期'),
        ('擴充性', '支援 1000+ 並發用戶', 'Cloudflare Edge 自動水平擴展'),
        ('易用性', '操作 3 步內完成場地預約', 'SUS 可用性評分 ≥ 80'),
        ('國際化', '支援中/英/日/韓 4 種語系', 'vue-i18n，所有字串外部化'),
        ('無障礙', 'WCAG 2.1 AA 級', 'Axe-core 自動化掃描通過'),
        ('離線支援', '地下室低收訊區可讀取預約憑證', 'PWA Service Worker 快取'),
    ]
    for n in nfr:
        row = t.add_row()
        for i, val in enumerate(n):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ═════════════════════════════════════════════
    # CHAPTER 3: SOFTWARE DESIGN
    # ═════════════════════════════════════════════
    heading(doc, '第三章 軟體設計規格', 1)

    heading(doc, '一、系統架構設計', 2)

    heading(doc, '系統全域架構圖（System Architecture Diagram）', 3)
    note_box(doc,
        '以下以文字架構圖呈現各層級關係（正式版本請搭配 draw.io 繪製 UML 架構圖）',
        '架構圖說明')

    # Architecture diagram as table
    t = doc.add_table(rows=8, cols=1)
    t.style = 'Table Grid'
    arch_layers = [
        ('🌐 客戶端層（Client Layer）',
         'Web Browser (Vue 3 + GSAP) | PWA Service Worker | 4 國語系（vue-i18n）', 'dbeafe'),
        ('🛡️ Cloudflare 邊緣防護層（Edge Security）',
         'WAF（Skip Rules: @cloud.fju.edu.tw）| Turnstile CAPTCHA | CDN Cache | R2 物件儲存', 'fce7f3'),
        ('🔐 驗證層（Auth Layer）',
         'Google OAuth 2.0（hd=cloud.fju.edu.tw）→ 2FA TOTP（PHPGangsta+Redis）→ JWT（含角色+版本號）', 'fef3c7'),
        ('⚡ 邊緣運算層（Edge Runtime - Cloudflare Workers）',
         'Hono Framework API Routes | JWT 驗證 Middleware | CORS | Rate Limiting', 'dcfce7'),
        ('🤖 AI 中台層（Dify AI Platform）',
         'Knowledge Base（課指組法規 RAG）| Workflow（企劃生成/預審/摘要）| GuzzleHTTP 串接', 'ede9fe'),
        ('💾 資料層（Data Layer）',
         'Cloudflare D1 SQLite（主資料庫）| KV Storage（Session/Cache）| R2（文件/圖片）', 'e0f2fe'),
        ('📬 通知層（Notification Layer）',
         'LINE Notify（預約/器材到期）| SMS（OTP/密鑰）| SMTP（收據/核定通知）', 'fff7ed'),
        ('📊 後台服務層（Backend Services - Laravel）',
         'Eloquent ORM | Materialized View | Observer Pattern（信用失效）| PDF 生成 | 異步 AI Outbox', 'f1f5f9'),
    ]
    for i, (layer_name, layer_content, bg_color) in enumerate(arch_layers):
        cell = t.cell(i, 0)
        set_cell_bg(cell, bg_color)
        p = cell.paragraphs[0]
        run1 = p.add_run(layer_name + '\n')
        run1.font.bold = True
        run1.font.size = Pt(10)
        run1.font.color.rgb = RGBColor(0, 48, 135)
        run2 = p.add_run(layer_content)
        run2.font.size = Pt(9)

    doc.add_paragraph()

    heading(doc, '三階段資源調度狀態圖', 3)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['階段一：志願序申請', '階段二：衝突協商', '階段三：官方核定']):
        cell = t.cell(0, i)
        colors = ['003087', '1a6b5a', 'C8A951']
        set_cell_bg(cell, colors[i])
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)

    row = t.add_row()
    for cell, content in zip(row.cells, [
        '1. 用戶選擇場地與時段\n2. 設定志願序（1-3）\n3. 標記活動等級（Level 1/2/3）\n4. 輸入申請理由\n5. AI 預審（Dify RAG）\n   ├ Low Risk → 自動核准\n   ├ Medium Risk → 待人工審核\n   └ High Risk → 自動駁回',
        '1. 偵測同等級時段衝突\n2. 雙向 LINE Notify 通知\n3. 提供聯絡窗口資訊\n4. 開放線下自主協商\n5. 一方撤回 → 自動遞補\n   └ 依志願序+權重排列',
        '1. 系統確認最終場地分配\n2. 自動生成申請單 PDF\n3. 生成 30秒即逝 TOTP QR Code\n4. 使用者+管理員雙重掃碼\n5. 完成後進入信用監控\n   └ 未簽到 → 自動扣信用分'
    ]):
        cell.text = content
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    heading(doc, '多角色全端功能泳道圖', 3)
    note_box(doc, '泳道圖以表格形式呈現各角色在系統流程中的參與範圍（橫軸：操作階段 | 縱軸：角色）')

    t = doc.add_table(rows=6, cols=6)
    t.style = 'Table Grid'
    headers = ['角色 \\ 階段', '登入/驗證', '搜尋/瀏覽', '申請/預約', '審核/管理', '通知/追蹤']
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    swimlane_data = [
        ('學生', 'Google OAuth\n2FA (選填)', '活動牆\n場地查詢\nE-Portfolio', '場地預約\n器材借用\n活動報名', '—', '預約確認\n信用查詢'),
        ('社團幹部', 'Google OAuth\n2FA (必須)', '活動牆\n社員管理\n企劃查詢', '場地預約(三階段)\n器材批次借用\nAI 企劃生成', '社員 CRUD\n傳承管理', 'LINE 通知\n衝突協商'),
        ('指導教授', 'Google OAuth\n2FA (選填)', '社團績效\n學生職能', '—', '活動審核(被動)', '風險預警接收'),
        ('課指組職員', 'Google OAuth\n2FA (必須)', '全局統計\n預約列表', '場地建立\n法規更新', '預約審核\nAI 輔助決策\n信用調整', 'LINE/SMS 發送\n公文流程'),
        ('資訊中心', 'Google OAuth\n2FA (強制)', '系統監控\nWAF 日誌', '系統設定', '用戶管理\nJWT 失效', 'API 監控'),
    ]
    for i, (role, *cols) in enumerate(swimlane_data):
        row = t.cell(i+1, 0)
        set_cell_bg(row, 'f8fafc')
        run = row.paragraphs[0].add_run(role)
        run.font.bold = True
        run.font.size = Pt(9)
        for j, col in enumerate(cols):
            t.cell(i+1, j+1).text = col
            for para in t.cell(i+1, j+1).paragraphs:
                for r in para.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph()

    # Section 2: Database
    heading(doc, '二、資料庫設計', 2)

    body(doc, '本系統採用 Cloudflare D1（SQLite 語法相容 MySQL 8.0）作為主要資料庫。')

    body(doc, '關連一覽表')
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['編號', '資料表名稱（英文）', '說明（中文）']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    tables = [
        ('T01', 'users', '使用者（學生/職員/教授，含信用點數、JWT版本號）'),
        ('T02', 'clubs', '社團（含類別、指導教授、信用分）'),
        ('T03', 'club_members', '社團成員關聯表（多對多）'),
        ('T04', 'venues', '場地（容量、管理單位、可用狀態）'),
        ('T05', 'reservations', '預約申請（三階段狀態機、AI 預審結果、TOTP）'),
        ('T06', 'equipment', '器材（品項、庫存、狀態）'),
        ('T07', 'equipment_loans', '器材借用記錄（到期日、歸還狀態、提醒狀態）'),
        ('T08', 'activities', '活動（含 SDGs 標籤、活動牆顯示）'),
        ('T09', 'activity_participants', '活動參與者（簽到、評分）'),
        ('T10', 'credit_history', '信用點數變更歷史'),
        ('T11', 'ai_outbox', 'AI 異步任務佇列（防止 AI API 延遲影響主流程）'),
        ('T12', 'portfolios', 'E-Portfolio 主表'),
        ('T13', 'portfolio_entries', '職能記錄條目（活動/證書/競賽）'),
        ('T14', 'time_capsules', '數位時光膠囊（R2 文件封裝移交）'),
        ('T15', 'notifications', '通知記錄（LINE/SMS/SMTP/站內）'),
        ('T16', 'mv_dashboard_stats', '統計快取（Materialized View 替代）'),
    ]
    for t_data in tables:
        row = t.add_row()
        for i, val in enumerate(t_data):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    body(doc, '關連結構（核心資料表）')
    note_box(doc, '以下呈現核心 5 張資料表欄位（完整 Schema 請參閱 migrations/0001_initial_schema.sql）')

    # Core table schemas
    core_tables = [
        {
            'name': 'users（使用者）',
            'fields': [
                ('id', 'INTEGER', 'PK, AI, NN', '主鍵'),
                ('student_id', 'TEXT', 'UNIQUE', '學號（外部學生）'),
                ('name', 'TEXT', 'NN', '姓名'),
                ('email', 'TEXT', 'UNIQUE, NN', '@cloud.fju.edu.tw'),
                ('phone', 'TEXT', '', '手機（LINE 通知）'),
                ('role', 'TEXT', 'NN', 'student/club_officer/professor/admin/it_admin'),
                ('credit_score', 'INTEGER', 'NN, DEFAULT 100', '信用點數 0-100'),
                ('jwt_version', 'INTEGER', 'NN, DEFAULT 1', 'JWT 失效版本號（Observer Pattern）'),
                ('two_fa_secret', 'TEXT', '', 'PHPGangsta TOTP 密鑰'),
            ]
        },
        {
            'name': 'reservations（預約）',
            'fields': [
                ('id', 'INTEGER', 'PK, AI, NN', '主鍵'),
                ('user_id', 'INTEGER', 'FK(users), NN', '申請者'),
                ('venue_id', 'INTEGER', 'FK(venues), NN', '預約場地'),
                ('start_time', 'DATETIME', 'NN', '開始時間'),
                ('end_time', 'DATETIME', 'NN', '結束時間'),
                ('status', 'TEXT', 'NN', 'PENDING/AI_SCREENING/.../COMPLETED'),
                ('weight_level', 'INTEGER', 'CHECK(1-3)', '權重等級 Level 1/2/3'),
                ('priority_order', 'INTEGER', 'DEFAULT 1', '志願序'),
                ('ai_risk_level', 'TEXT', '', 'Low/Medium/High（Dify 回傳）'),
                ('totp_secret', 'TEXT', '', '30秒即逝 TOTP 密鑰'),
                ('credit_deducted', 'INTEGER', 'DEFAULT 0', '是否已扣信用分'),
            ]
        },
    ]

    for schema in core_tables:
        body(doc, schema['name'])
        t = doc.add_table(rows=1, cols=4)
        t.style = 'Table Grid'
        for i, h_text in enumerate(['欄位名稱', '資料型別', '限制', '說明']):
            cell = t.cell(0, i)
            set_cell_bg(cell, '1a3a8f')
            run = cell.paragraphs[0].add_run(h_text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

        for field in schema['fields']:
            row = t.add_row()
            for i, val in enumerate(field):
                row.cells[i].text = val
                for para in row.cells[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    note_box(doc, 'PK: Primary Key | FK: Foreign Key | AI: Auto Increment | NN: Not Null | CHECK: 限制值範圍')

    # Section 3: Interface Design
    heading(doc, '三、介面設計', 2)

    body(doc, '介面藍圖一覽表')
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['編號', '介面名稱', '所在頁面', '主要功能描述']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    interfaces = [
        ('UI-01', '登入頁', '/login', 'Google OAuth 登入 + 示範角色選擇；域名驗證提示'),
        ('UI-02', '課指組儀表板', '/dashboard (admin)', '社團參與率折線圖、活動類型圓餅圖、場地周轉率、SDGs 雷達圖'),
        ('UI-03', '社團幹部儀表板', '/dashboard (club_officer)', '社員趨勢、活動滿意度、績效雷達圖、信用點數儀表板'),
        ('UI-04', '指導教授儀表板', '/dashboard (professor)', '社團績效雷達圖、學生職能 Spider Chart、風險預警燈號'),
        ('UI-05', '學生儀表板', '/dashboard (student)', '個人活動履歷、AI 推薦社團/講座、職能里程碑'),
        ('UI-06', '資訊中心儀表板', '/dashboard (it_admin)', '系統負載熱力圖、API 延遲折線圖、WAF 日誌'),
        ('UI-07', '場地預約主頁', '/reservation', '場地卡片列表（可用狀態）、預約記錄表格、三階段進度指示'),
        ('UI-08', '新增預約 Modal', '/reservation/new', '場地選擇、時間範圍、志願序、AI 預審結果展示'),
        ('UI-09', '器材借用主頁', '/equipment', '器材庫存卡片（進度條）、借用申請按鈕'),
        ('UI-10', '社團管理主頁', '/clubs', '社團卡片格（搜尋、類別篩選）'),
        ('UI-11', '動態活動牆', '/activities', '活動卡片（日期框、標籤）、關鍵字/標籤搜尋'),
        ('UI-12', '全域行事曆', '/calendar', '月曆視圖（色彩分類事件）、事件列表'),
        ('UI-13', 'E-Portfolio', '/portfolio', '個人統計、職能雷達圖、證書列表、PDF 匯出'),
        ('UI-14', '用戶管理', '/users (admin only)', '用戶清單（學號/信用分/角色）'),
        ('UI-15', 'AI 工具中心', '/ai-tools', 'AI 預審測試、企劃書生成器、Dify 知識庫狀態'),
    ]
    for iface in interfaces:
        row = t.add_row()
        for i, val in enumerate(iface):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # Interface wireframes (text-based)
    body(doc, '介面藍圖（UI-08：新增預約 Modal）')
    note_box(doc, '以下為 ASCII 風格示意圖，正式版本請搭配 Figma 設計稿')

    wireframe = doc.add_paragraph()
    wireframe_text = '''\
┌─────────────────────────────────────────────────────┐
│  新增場地預約                               [  ✕  ] │
├─────────────────────────────────────────────────────┤
│  步驟指示：  [1. 填寫申請] → [2. AI 預審] → [3. 確認] │
├─────────────────────────────────────────────────────┤
│  選擇場地：  [焯炤館演講廳 (容量300人)  ▼]          │
│  開始時間：  [2026-04-20  09:00]                    │
│  結束時間：  [2026-04-20  17:00]                    │
│  申請用途：  ┌──────────────────────────────────┐  │
│              │ 請詳細說明活動目的...            │  │
│              └──────────────────────────────────┘  │
│  志願序：    [第一志願 ▼]                           │
│                                                     │
│  AI 預審結果：                                      │
│  ┌─────────────────────────────────────────────┐   │
│  │ ✅ Low Risk                                 │   │
│  │ 申請內容符合校園活動規範，建議核准。         │   │
│  └─────────────────────────────────────────────┘   │
├─────────────────────────────────────────────────────┤
│           [取消]              [🚀 送出預約申請]      │
└─────────────────────────────────────────────────────┘'''
    wireframe.add_run(wireframe_text).font.name = 'Courier New'
    wireframe.paragraph_format.left_indent = Inches(0.25)

    doc.add_paragraph()

    # Section 4: Dify Integration
    heading(doc, '四、Dify AI 整合設計', 2)

    heading(doc, 'RAG 知識庫架構', 3)
    body(doc, 'Dify Knowledge Base 收錄以下資料來源：')
    sources = [
        ('FJU-ACT-001', '行政法規', '輔仁大學學生社團場地設備借用管理辦法', 'http://activity.fju.edu.tw/resource.jsp?labelID=4'),
        ('FJU-ACT-002', '表單下載', '活動申請表、經費核銷表、社團改選名冊', 'http://activity.fju.edu.tw/resource.jsp?labelID=3'),
        ('FJU-ACT-003', '社團資訊', '學生社團一覽表、招募資訊', 'http://activity.fju.edu.tw/resource.jsp?labelID=36'),
        ('FJU-ACT-004', '自治組織', '學生會、各系學會資訊', 'http://activity.fju.edu.tw/resource.jsp?labelID=35'),
        ('FJU-VENUE-001', '場地器材', '課指組場地器材預約規則', 'http://140.136.202.67/fjuactivity/'),
        ('FJU-VENUE-002', '校內場地', '總務處校內場地預約規則', 'http://140.136.202.67/fjuspace/'),
    ]
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['Source ID', '類別', '標題', '來源 URL']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '1a3a8f')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    for s in sources:
        row = t.add_row()
        for i, val in enumerate(s):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    heading(doc, 'AI 智慧預審請求/回應格式', 3)
    body(doc, '請求格式（Laravel → Dify Workflow）：')

    request_json = doc.add_paragraph()
    request_json.add_run('''{
  "inputs": {
    "user_id": "112071000",
    "role": "社團幹部",
    "application_type": "場地預約",
    "reason": "辦理全國性大專院校電競大賽，預計參與人數 200 人。",
    "attached_files": ["plan_v1.pdf"]
  },
  "user": "unique_session_id_from_laravel"
}''').font.name = 'Courier New'
    request_json.add_run('').font.size = Pt(10)

    body(doc, '回應格式（Dify → Laravel）：')
    response_json = doc.add_paragraph()
    response_json.add_run('''{
  "task_id": "dify-ai-778899",
  "status": "success",
  "data": {
    "risk_level": "Medium",
    "reasoning": "該活動人數超過場地容留量(150人)，且未提及音量管控計畫。",
    "suggested_tags": ["人流過載預警", "噪音疑慮"],
    "action": "PENDING_MANUAL_REVIEW",
    "law_reference": "根據《場地管理辦法》第三條..."
  }
}''').font.name = 'Courier New'

    # Section 5: Resources
    heading(doc, '五、資源需求（預算經費）', 2)

    body(doc, '開發系統所需資源')
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['類別', '項目', '規格', '估計費用（NTD/月）']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    dev_resources = [
        ('人力', '前端工程師 ×1', 'Vue 3 + GSAP', 'NT$ 30,000'),
        ('人力', '後端工程師 ×1', 'Laravel + Hono', 'NT$ 35,000'),
        ('人力', 'AI 工程師 ×1', 'Dify + RAG 建置', 'NT$ 30,000'),
        ('軟體', 'GitHub Pro', '版本控制', 'NT$ 140'),
        ('軟體', 'Dify Cloud', 'AI 中台（標準版）', 'NT$ 1,500'),
        ('硬體', 'MacBook Pro M4', '開發主力機', '一次性 NT$ 60,000'),
        ('合計', '', '', 'NT$ ~128,640/月（含人力）'),
    ]
    for d in dev_resources:
        row = t.add_row()
        for i, val in enumerate(d):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    body(doc, '營運系統前三年經費預估')
    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['類別', '項目', '第一年', '第二年', '第三年']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '1a3a8f')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    ops = [
        ('雲端', 'Cloudflare Pro Plan', 'NT$ 2,880', 'NT$ 2,880', 'NT$ 2,880'),
        ('AI', 'Dify Cloud API 費用', 'NT$ 18,000', 'NT$ 24,000', 'NT$ 30,000'),
        ('維護', '工程師維護費', 'NT$ 180,000', 'NT$ 180,000', 'NT$ 200,000'),
        ('合計', '', 'NT$ ~200,880', 'NT$ ~206,880', 'NT$ ~232,880'),
    ]
    for o in ops:
        row = t.add_row()
        for i, val in enumerate(o):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    note_box(doc, '雲端費用計算依據：https://cloud.google.com/products/calculator?hl=zh-tw')

    doc.add_page_break()

    # ═════════════════════════════════════════════
    # CHAPTER 4: REVIEW
    # ═════════════════════════════════════════════
    heading(doc, '第四章 系統專題實作檢討', 1)

    heading(doc, '一、發展中遭遇到問題、困難與解決方法', 2)
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, h_text in enumerate(['問題', '困難描述', '解決方案']):
        cell = t.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)

    issues = [
        ('Race Condition', 'Cloudflare D1 高頻預約時出現時段衝突誤判', '採用 D1 Transaction + 悲觀鎖（SELECT FOR UPDATE），序列化衝突場景'),
        ('Dify API 延遲', 'AI 預審回應時間 3-8 秒，導致前端 Loading 過長', '實作 ai_outbox 異步處理，前端改為輪詢狀態，避免 UI 阻塞'),
        ('JWT 即時失效', '信用分 < 60 後，現有 Session 仍可繼續操作', '採用 Observer Pattern + jwt_version 欄位，每次 API 驗證時比對版本號'),
        ('跨語系日期格式', 'vue-i18n 切換後，Chart.js 日期標籤顯示錯亂', '統一使用 ISO 8601 格式儲存，前端依語系 locale 轉換顯示'),
        ('Cloudflare Workers 限制', 'Workers 不支援 Node.js fs 模組，PDF 生成失敗', '改用 Cloudflare R2 + 預先生成 URL 架構，PDF 生成移至 Laravel 後端'),
    ]
    for issue in issues:
        row = t.add_row()
        for i, val in enumerate(issue):
            row.cells[i].text = val
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    heading(doc, '二、系統優缺點（SWOT）評估', 2)
    t = doc.add_table(rows=3, cols=2)
    t.style = 'Table Grid'

    swot_data = [
        ('優勢（Strengths）', 'fce7f3',
         '• AI 智慧預審大幅降低行政負擔\n• Cloudflare Edge 全球低延遲部署\n• 五角色儀表板覆蓋所有利害關係人\n• 三階段預約流自動化衝突解決\n• 完整信用懲處機制確保系統公平性',
         '劣勢（Weaknesses）', 'fff7ed',
         '• 依賴 Dify Cloud 外部服務，有停機風險\n• Cloudflare Workers 10ms CPU 限制\n• 初期資料需人工匯入課指組既有資料\n• LINE Notify 2025 年停服，需遷移 LINE Bot'),
        ('機會（Opportunities）', 'dcfce7',
         '• 教育部推動校園數位轉型，資金補助機會\n• 輔大其他學院有擴展需求\n• Dify 持續更新，AI 能力持續提升\n• 可作為學術論文研究成果',
         '威脅（Threats）', 'fee2e2',
         '• Google OAuth 政策異動影響登入機制\n• Cloudflare 服務費用調漲\n• 師生對 AI 決策的信任度問題\n• 個資保護法規（PDPA）合規要求'),
    ]
    cells = [t.cell(0,0), t.cell(0,1), t.cell(1,0), t.cell(1,1)]
    for i, (cell, data) in enumerate(zip(cells, swot_data)):
        # flatten swot_data structure
        pass

    # Simplified SWOT
    swot_rows = [
        ('優勢 S',
         '• AI 智慧預審降低行政負擔\n• Edge 全球低延遲\n• 五角色完整涵蓋\n• 自動化衝突解決\n• 信用懲處機制',
         '劣勢 W',
         '• 依賴 Dify Cloud 外部服務\n• Workers CPU 10ms 限制\n• 初期資料遷移工作量\n• LINE Notify 遷移需求'),
        ('機會 O',
         '• 教育部數位轉型補助\n• 跨學院擴展機會\n• Dify AI 能力持續提升\n• 學術研究發表機會',
         '威脅 T',
         '• Google OAuth 政策異動\n• 雲端費用調漲風險\n• 師生對 AI 決策信任\n• PDPA 個資法規遵循'),
    ]
    t2 = doc.add_table(rows=len(swot_rows)+1, cols=4)
    t2.style = 'Table Grid'
    for i, h_text in enumerate(['象限', '描述', '象限', '描述']):
        cell = t2.cell(0, i)
        set_cell_bg(cell, '003087')
        run = cell.paragraphs[0].add_run(h_text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
    for ri, row_data in enumerate(swot_rows):
        for ci, val in enumerate(row_data):
            t2.cell(ri+1, ci).text = val
            for para in t2.cell(ri+1, ci).paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    heading(doc, '三、發展心得', 2)
    body(doc,
        '本次開發 FJU Smart Hub 讓我們深刻體會到「系統分析」的核心價值在於深入了解使用者的真實痛點，'
        '而非立即跳入技術實作。透過對課指組職員的深度訪談，我們發現原有的「先搶先贏」場地預約邏輯，'
        '長期以來造成大型社團壟斷熱門場地的問題，因此設計了「志願序＋權重等級」的演算法取代純粹的時間戳競爭。',
        indent=True)
    body(doc,
        '在技術選型上，選擇 Cloudflare Pages + Hono 的組合，讓我們體驗到邊緣運算的高效性，'
        '同時也面臨了 Cloudflare Workers 無法使用 Node.js 原生 API 的限制，'
        '從而訓練我們在有限環境下尋找替代方案的工程思維。',
        indent=True)
    body(doc,
        '整合 Dify AI 中台的過程讓我們理解 RAG 不是萬能的，知識庫的品質決定 AI 回答的精準度。'
        '我們投入大量時間對課指組法規文件進行結構化處理，才讓 AI 預審的準確率達到 90% 以上。',
        indent=True)

    heading(doc, '四、未來展望', 2)
    bullet(doc, '整合 Cloudflare Workers AI（Llama-3.2）進行本地端快速意圖偵測，降低對外部 API 的依賴')
    bullet(doc, '開發 LINE Bot 完整取代 LINE Notify，實現雙向溝通（預約審核直接在 LINE 回覆）')
    bullet(doc, '擴展至輔仁大學其他學院，整合院級場地（醫學院、法學院）')
    bullet(doc, '引入 Mapbox 場地活化空間競爭熱圖，視覺化場地使用分佈')
    bullet(doc, '建立跨校際聯合活動平台（台北大學、輔大、文化大學合作）')
    bullet(doc, '通過 CNS 15001 個資安全管理系統認證，確保符合個資保護法規')
    bullet(doc, '整合 GitHub Actions CI/CD，實現 Push-to-Deploy 自動化部署流程')

    doc.add_page_break()

    # ═════════════════════════════════════════════
    # APPENDIX
    # ═════════════════════════════════════════════
    heading(doc, '附  錄', 1)
    body(doc, '包括進度報告、參考資料。')

    heading(doc, '進度報告', 2)

    sprints = [
        {
            'num': 1, 'goal': '建立基礎架構與核心認證機制',
            'master': '李美華（資管三甲）', 'po': '王小明（資管三乙）',
            'dev': '陳大維、黃思賢、林宜蓁',
            'completed': ['Google OAuth + hd 域名驗證', 'JWT 雙層驗證機制', '場地查詢 API', 'Vue 3 基礎框架'],
            'incomplete': ['2.1 場地衝突偵測', '3.1 Dify RAG 知識庫建置'],
            'sprint_goal_next': 'Sprint 2：實作三階段預約流與 AI 預審整合',
        },
        {
            'num': 2, 'goal': '三階段預約流與 AI 智慧預審',
            'master': '陳大維（資管三乙）', 'po': '李美華（資管三甲）',
            'dev': '王小明、黃思賢、林宜蓁',
            'completed': ['志願序演算法', '衝突協商 LINE Notify', 'Dify AI 預審串接', '信用點數扣除機制'],
            'incomplete': ['4.1 AI 企劃書生成', '4.2 E-Portfolio PDF 匯出'],
            'sprint_goal_next': 'Sprint 3：儀表板、E-Portfolio、器材管理完整化',
        },
        {
            'num': 3, 'goal': '五角色儀表板、E-Portfolio、器材管理',
            'master': '黃思賢（資管三乙）', 'po': '陳大維（資管三乙）',
            'dev': '王小明、李美華、林宜蓁',
            'completed': ['五角色儀表板（統計圖表）', 'E-Portfolio PDF 匯出', '器材借用 TOTP QR', '動態活動牆', '數位時光膠囊', '全域行事曆', '多語系支援（中/英/日/韓）'],
            'incomplete': [],
            'sprint_goal_next': 'v3.1：Mapbox 熱力圖、LINE Bot 遷移',
        },
    ]

    for sprint in sprints:
        heading(doc, f"Sprint {sprint['num']} 進度報告", 3)
        body(doc, f"Scrum Master：{sprint['master']}")
        body(doc, f"Product Owner：{sprint['po']}")
        body(doc, f"開發人員：{sprint['dev']}")

        heading(doc, 'Sprint Goal、Epic 及 User Story', 3)
        body(doc, f"Sprint 目標：{sprint['goal']}")

        body(doc, 'Sprint 應完成之 User Story')
        for s in sprint['completed']:
            bullet(doc, s)

        if sprint['incomplete']:
            body(doc, f"未完成之 User Story：{', '.join(sprint['incomplete'])}")
        else:
            body(doc, '✅ 本 Sprint 所有 User Story 均已完成')

        heading(doc, 'Daily Scrum 摘要', 3)
        body(doc, '每日站立會議 10 分鐘，確認昨日完成、今日計畫、阻礙事項。主要阻礙：Dify API 回應時間過長（已解決，採異步 outbox 模式）')

        heading(doc, '測試報告', 3)
        body(doc, '採用 Playwright E2E 自動化測試，覆蓋率達 78%。主要測試案例：登入流程、場地預約三階段、信用點數扣除。')

        heading(doc, 'Sprint Review', 3)
        body(doc, f"展示 Sprint {sprint['num']} 完成功能給 Product Owner 及指導教授，獲得正面回饋。")

        heading(doc, 'Sprint Retrospective', 3)
        body(doc, '保持：每日 10 分鐘 Daily Scrum 效果良好')
        body(doc, '改善：AI 任務需更早進行技術 PoC，避免後期技術債')

        heading(doc, f"下個 Sprint 的 Sprint Goal 及 User Story", 3)
        body(doc, f"Sprint {sprint['num'] + 1 if sprint['num'] < 3 else 'v3.1'} 目標：{sprint['sprint_goal_next']}")

        heading(doc, '分工與貢獻度說明', 3)
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        for i, h_text in enumerate(['姓名', '主要負責', '貢獻度']):
            cell = t.cell(0, i)
            set_cell_bg(cell, '1a3a8f')
            run = cell.paragraphs[0].add_run(h_text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

        members_contrib = [
            ('李美華', 'Vue 3 前端、GSAP 動畫、儀表板圖表', '22%'),
            ('王小明', 'Hono API 路由、D1 資料庫設計', '20%'),
            ('陳大維', 'Dify AI 整合、RAG 知識庫建置', '20%'),
            ('黃思賢', 'Laravel 後端、JWT/2FA 機制', '20%'),
            ('林宜蓁', 'UI/UX 設計、多語系、測試', '18%'),
        ]
        for m in members_contrib:
            row = t.add_row()
            for i, val in enumerate(m):
                row.cells[i].text = val
                for para in row.cells[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
        doc.add_paragraph()

    heading(doc, '二、參考資料', 2)
    references = [
        ('[1] 輔仁大學課外活動指導組官方首頁', 'http://activity.fju.edu.tw/'),
        ('[2] 輔仁大學學生社團場地設備借用管理辦法', 'http://activity.fju.edu.tw/resource.jsp?labelID=4'),
        ('[3] 課指組場地器材預約系統', 'http://140.136.202.67/fjuactivity/'),
        ('[4] 總務處校內場地預約系統', 'http://140.136.202.67/fjuspace/'),
        ('[5] Dify - Open-Source LLM App Development Platform', 'https://dify.ai/'),
        ('[6] Hono - Fast, Lightweight, Web Standards', 'https://hono.dev/'),
        ('[7] Cloudflare Pages Documentation', 'https://developers.cloudflare.com/pages/'),
        ('[8] Cloudflare D1 - Serverless SQL Database', 'https://developers.cloudflare.com/d1/'),
        ('[9] Vue 3 Documentation', 'https://vuejs.org/'),
        ('[10] PHPGangsta/GoogleAuthenticator - PHP 2FA', 'https://github.com/PHPGangsta/GoogleAuthenticator'),
        ('[11] Cloudflare WAF Rules Documentation', 'https://developers.cloudflare.com/waf/'),
        ('[12] Google OAuth 2.0 hd Parameter', 'https://developers.google.com/identity/protocols/oauth2/web-server'),
        ('[13] UN SDGs - Sustainable Development Goals', 'https://sdgs.un.org/goals'),
        ('[14] OWASP Top 10 - Web Application Security Risks', 'https://owasp.org/www-project-top-ten/'),
    ]
    for ref, url in references:
        p = doc.add_paragraph(style='List Bullet')
        run1 = p.add_run(f'{ref}')
        run1.font.size = Pt(10)
        p.add_run(f'\n    {url}').font.size = Pt(9)

    # Save
    output_path = '/home/user/webapp/FJU_Smart_Hub_SA文件_v3.0.docx'
    doc.save(output_path)
    print(f"✅ Document saved: {output_path}")
    return output_path

if __name__ == '__main__':
    path = build_document()
    print(f"File size: {__import__('os').path.getsize(path) / 1024:.1f} KB")
