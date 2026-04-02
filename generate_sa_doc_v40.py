#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FJU Smart Hub v4.0 - System Analysis Specification Document Generator
Generates Word document with auto Table of Contents
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colours ──
BLUE      = RGBColor(0x00, 0x31, 0x53)   # Holy Mother Blue
GOLD      = RGBColor(0xDA, 0xA5, 0x20)   # Vatican Gold
GREEN     = RGBColor(0x00, 0x80, 0x00)   # Jiahe Green
RED       = RGBColor(0xFF, 0x00, 0x00)   # Alert Red
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY      = RGBColor(0x64, 0x74, 0x8B)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_toc(doc):
    """Add Word auto-generated TOC field"""
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles['Normal']
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-4" \\h \\z \\u'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)

def add_heading(doc, text, level=1, color=None):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if color:
            run.font.color.rgb = color
        else:
            run.font.color.rgb = BLUE if level <= 2 else GRAY
    return heading

def add_styled_para(doc, text, bold=False, color=None, size=None, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if size: run.font.size = Pt(size)
    return para

def add_table_header(table, headers, bg='003153', text_color=WHITE):
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = text_color
                run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_code_block(doc, code):
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    return para

def add_mermaid_block(doc, title, code):
    doc.add_paragraph()
    p = add_styled_para(doc, f'【流程圖】{title}', bold=True, color=GOLD)
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    para.paragraph_format.left_indent = Inches(0.3)

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ─── COVER PAGE ───
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('FJU Smart Hub v4.0')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BLUE

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run('系統分析規格書')
    run2.font.size = Pt(20)
    run2.font.color.rgb = GOLD

    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'輔仁大學校園智慧管理系統\n版本：v4.0   日期：{datetime.date.today()}\n作者：FJU Smart Hub 開發團隊').font.size = Pt(12)

    doc.add_page_break()

    # ─── AUTO TABLE OF CONTENTS ───
    add_heading(doc, '目錄', level=1)
    add_styled_para(doc, '（請按 F9 更新目錄）', color=GRAY, italic=True)
    doc.add_paragraph()
    add_toc(doc)
    doc.add_page_break()

    # ════════════════════════════════════════════
    # 1. 系統概述
    # ════════════════════════════════════════════
    add_heading(doc, '1. 系統概述', level=1)

    add_heading(doc, '1.1 系統背景與目標', level=2)
    add_styled_para(doc, 'FJU Smart Hub v4.0 是輔仁大學校園智慧管理系統，旨在整合場地預約、社團管理、AI輔助協商、信用積分追蹤及無障礙設施導覽等核心功能。本版本為全面升級版本，引入 Leaflet.js 互動式校園地圖、玻璃態行事曆面板、三/六分鐘 AI 仲裁協商機制及完整的已讀回條追蹤系統。')

    add_heading(doc, '1.2 視覺設計規範（60-30-10）', level=2)
    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = 'Table Grid'
    add_table_header(tbl, ['角色', '顏色代號', '應用位置'])
    data = [
        ('主色（60%）', 'Angel White #FFFFFF', '背景、卡片底色'),
        ('次色（30%）', 'Holy Mother Blue #003153', '導覽列、側欄標題、CTA 按鈕'),
        ('強調（10%）', 'Vatican Gold #DAA520', '搜尋高亮、信用點數條、無障礙圖示'),
        ('狀態色', 'Jiahe Green #008000 / Alert Red #FF0000', '正常/完成 / 維護中/逾期'),
    ]
    for i, row_data in enumerate(data):
        row = tbl.rows[i+1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    add_heading(doc, '1.3 Bento Grid 桌面佈局', level=2)
    add_styled_para(doc, '■ Header（頂部）：左側大學 Logo、中央浮動 Google Maps 風格搜尋列、右側通知鈴鐺（紅點）＋用戶頭像', bold=False)
    add_styled_para(doc, '■ 左側欄（≈30%）：信用積分儀表板、各學院訊息卡、待協商清單、功能分類快捷鍵')
    add_styled_para(doc, '■ 右主區（≈70%）：Base Layer（Leaflet 全螢幕地圖）+ Floating Layer（玻璃態行事曆面板，40%寬，GSAP滑入）')

    add_heading(doc, '1.4 技術架構總覽', level=2)
    tbl2 = doc.add_table(rows=8, cols=2)
    tbl2.style = 'Table Grid'
    add_table_header(tbl2, ['層級', '技術選型'])
    tech = [
        ('前端', 'HTML5 + CSS3 (Bento Grid) + JavaScript ES13+'),
        ('地圖', 'Leaflet.js 1.9.4 + SVG Overlay + Accessibility Markers'),
        ('動畫', 'GSAP 3.12 (行事曆滑入、Landing Page 動效)'),
        ('圖表', 'Chart.js 4.4 (折線、圓餅、雷達、柱狀)'),
        ('後端（模擬）', 'Hono + TypeScript on Cloudflare Workers/Pages'),
        ('資料庫', 'Cloudflare D1 (SQLite) + KV + R2（生產環境設計）'),
        ('認證', '@cloud.fju.edu.tw + 2FA (TOTP/SMS) + JWT HttpOnly Cookie'),
    ]
    for i, row_data in enumerate(tech):
        row = tbl2.rows[i+1]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        for c in row.cells:
            c.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 2. 系統架構圖
    # ════════════════════════════════════════════
    add_heading(doc, '2. 系統架構圖', level=1)

    add_heading(doc, '2.1 整體架構 (Mermaid)', level=2)
    add_mermaid_block(doc, '整體系統架構', '''
graph TB
  subgraph Frontend ["前端層 (Browser)"]
    LP[Landing Page<br/>GSAP動效]
    LGN[Login Page<br/>Role Selection]
    HEADER[Header<br/>Logo + Search + Bell + Avatar]
    SIDEBAR[左側欄 30%<br/>Credit | News | Nego | Shortcuts]
    MAP[Leaflet.js Map<br/>Base Layer]
    CAL[玻璃態行事曆<br/>Floating Layer]
    PAGES[功能頁面<br/>Dashboard / Reservation / etc]
  end
  subgraph Backend ["後端層 (Hono Workers)"]
    AUTH[Auth API<br/>/api/auth/*]
    VENUE[Venue API<br/>/api/venues/*]
    RES[Reservation API<br/>/api/reservations/*]
    AI[AI APIs<br/>/api/ai/*]
    MAP_API[Map API<br/>/api/map/*]
    CREDIT[Credit API<br/>/api/credit/*]
    NOTIF[Notification API<br/>/api/notifications/*]
    CONFLICT[Conflict API<br/>/api/conflicts/*]
  end
  subgraph DB ["資料層"]
    D1[(D1 SQLite<br/>主要資料庫)]
    KV[(KV Store<br/>Session/Cache)]
    R2[(R2 Storage<br/>檔案/PDF)]
  end
  Frontend --> Backend
  Backend --> DB
  AUTH --> KV
  RES --> D1
  CREDIT --> D1
  NOTIF --> D1
  MAP_API --> D1
  CONFLICT --> D1
''')

    add_heading(doc, '2.2 前端組件樹', level=2)
    add_mermaid_block(doc, '前端組件架構', '''
graph TD
  A[app.js Entry] --> B[renderLanding]
  A --> C[renderLogin]
  A --> D[initApp]
  D --> E[buildAppHTML]
  E --> F[Header Component]
  E --> G[Sidebar Component]
  E --> H[Main Area]
  G --> G1[Credit Dashboard]
  G --> G2[College News Cards]
  G --> G3[Negotiation List]
  G --> G4[Function Shortcuts]
  H --> H1[Leaflet Map Layer]
  H --> H2[Calendar Panel - glassmorphism]
  H --> H3[Content Pages]
  H2 --> H21[Calendar Grid]
  H2 --> H22[Event List]
  H2 --> H23[Go to Cal Management]
  H3 --> H31[Dashboard]
  H3 --> H32[Reservation 3-phase]
  H3 --> H33[Equipment]
  H3 --> H34[Clubs]
  H3 --> H35[Activities]
  H3 --> H36[AI Tools]
  H3 --> H37[Portfolio]
''')

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 3. 功能模組說明
    # ════════════════════════════════════════════
    add_heading(doc, '3. 功能模組說明', level=1)

    add_heading(doc, '3.1 動畫引導頁面（Landing Page）', level=2)
    add_styled_para(doc, '【設計原則】Landing Page 不是登入畫面，而是展示系統功能的動畫介紹頁面。')
    add_styled_para(doc, '■ 視覺元素：深藍漸層背景、金色粒子動效、大學 Logo 光環脈衝動畫')
    add_styled_para(doc, '■ 功能展示卡：互動式校園地圖、AI衝突協商、玻璃行事曆、信用積分制、已讀回條 × 5張')
    add_styled_para(doc, '■ 進入按鈕：渐变金色按鈕，點擊後 fade out → 顯示登入頁面')
    add_styled_para(doc, '■ 技術：CSS @keyframes + JavaScript particle generation（20顆浮動粒子）')

    add_heading(doc, '3.2 認證系統（Auth）', level=2)
    add_styled_para(doc, '■ 僅限 @cloud.fju.edu.tw 帳號')
    add_styled_para(doc, '■ 2FA：TOTP（Google Authenticator）+ SMS 備援')
    add_styled_para(doc, '■ JWT 存放在 HttpOnly Cookie（防 XSS）')
    add_styled_para(doc, '■ 角色：student / club_officer / professor / admin / it_admin')
    add_styled_para(doc, '■ JWT version 欄位：用於強制登出（Observer Pattern）')

    add_heading(doc, '3.3 Leaflet.js 互動式校園地圖', level=2)
    add_styled_para(doc, '■ Base Layer：OpenStreetMap 底圖，中心點 25.0330, 121.4330，縮放 z=17')
    add_styled_para(doc, '■ 建築標記：共 10 棟，圓形圖示（藍色=正常, 紅色=維護中）')
    add_styled_para(doc, '■ 無障礙標記：金點（坡道）、綠方（電梯）、紫圓（廁所）、橙菱（停車位）')
    add_styled_para(doc, '■ 篩選按鈕：全部 / 無障礙設施 / 可預約 / 維護中')
    add_styled_para(doc, '■ 側欄 click → flyTo：點擊側欄快捷鍵時，地圖執行 flyTo(lat,lng, zoom=18)')
    add_styled_para(doc, '■ building_id AJAX：點擊建築圖示發送 GET /api/map/building/:id，即時獲取維護狀態（來自 map_elements 資料表）')

    add_heading(doc, '3.4 玻璃態行事曆面板（Glassmorphism Calendar）', level=2)
    add_styled_para(doc, '■ 設計：backdrop-filter:blur(20px)+saturate(180%)，半透明玻璃效果，40%寬、右側對齊')
    add_styled_para(doc, '■ 動畫：translateX 從 105% → 0（GSAP 或 CSS transition 0.4s cubic-bezier）')
    add_styled_para(doc, '■ 功能：月份導航、事件點（綠/金/紅點）、點擊日期顯示詳情')
    add_styled_para(doc, '■ 衝突日期：點擊衝突狀態事件 → 彈出協商對話框')
    add_styled_para(doc, '■「開啟行事曆管理頁面」按鈕：跳轉至獨立行事曆管理頁（全月格狀視圖）')

    add_heading(doc, '3.5 AI 衝突協商機制（3/6 分鐘規則）', level=2)
    add_mermaid_block(doc, 'AI 衝突協商流程', '''
flowchart TD
  A([場地申請衝突偵測]) --> B[開啟協商室]
  B --> C[開始計時 00:00]
  C --> D{用戶發送訊息?}
  D -->|是| E[重置沉默計時器]
  D -->|否 30秒| F[顯示沉默警告]
  E --> G{已達 3 分鐘?}
  F --> G
  G -->|否| D
  G -->|是| H[GPT-4 AI 介入<br/>提供 3 項具體建議]
  H --> I{用戶接受建議?}
  I -->|是| J[雙方達成協議]
  I -->|否| K{已達 6 分鐘?}
  K -->|否| I
  K -->|是| L[強制處置]
  L --> M[關閉對話框]
  L --> N[扣除 10 信用點數]
  L --> O[紅燈閃爍動效]
  L --> P[記錄 credit_logs]
  J --> Q[更新 conflicts 表]
  J --> R[加 2 信用點數]
  J --> S[跳轉行事曆管理頁]
''')
    add_styled_para(doc, '■ 3分鐘規則：計時器到達 180 秒，若語意過濾判定雙方沉默或內容無關，GPT-4 自動介入提供 3 個具體建議')
    add_styled_para(doc, '■ 6分鐘規則：計時器到達 360 秒仍無共識，強制關閉對話框、扣除 10 信用點數、顯示紅燈閃爍動效、記錄 credit_logs 表')
    add_styled_para(doc, '■「協商完成，更新行事曆」：更新 conflicts 表 status=resolved，加 2 信用點數，跳轉行事曆管理頁面')

    add_heading(doc, '3.6 已讀回條追蹤（Read Receipt / track.php）', level=2)
    add_styled_para(doc, '■ 每則通知附帶唯一 token（UUID格式）')
    add_styled_para(doc, '■ 用戶點擊通知 → POST /api/notifications/track { token, timestamp } → 記錄 notification_logs 表')
    add_styled_para(doc, '■ 重要通知（is_mandatory=true）：觸發 10 秒強制 modal overlay，倒數結束後啟用確認按鈕')
    add_styled_para(doc, '■ 紀錄欄位：token, read_at, ip_address, user_agent, confirmed')
    add_mermaid_block(doc, '已讀回條流程', '''
sequenceDiagram
  participant U as 用戶
  participant FE as 前端 JS
  participant BE as Hono Backend
  participant DB as notification_logs
  U->>FE: 點擊通知
  FE->>FE: 渲染通知內容
  alt 重要通知
    FE->>FE: 顯示 10s 強制 Modal
    FE->>U: 倒數 10 秒
    U->>FE: 點擊「已閱讀確認」
  end
  FE->>BE: POST /api/notifications/track { token, timestamp }
  BE->>DB: INSERT notification_logs (token, read_at, ip_address, user_agent)
  DB-->>BE: Success
  BE-->>FE: { success:true, ip, tracked_at }
  FE->>FE: 標記已讀 UI 更新
''')

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 4. ER 圖（完整資料庫設計）
    # ════════════════════════════════════════════
    add_heading(doc, '4. 資料庫設計（ER 圖）', level=1)

    add_heading(doc, '4.1 完整 ER 圖（Mermaid）', level=2)
    add_mermaid_block(doc, '完整 ER 圖', '''
erDiagram
  users {
    int id PK
    string student_id UK
    string name
    string email UK
    string role
    string department
    int credit_score
    int is_active
    string google_id
    string two_fa_secret
    int jwt_version
    datetime last_login_at
    datetime created_at
  }
  clubs {
    int id PK
    string name
    string category
    int advisor_id FK
    int credit_score
    string status
    datetime created_at
  }
  venues {
    int id PK
    string name
    int capacity
    string type
    string managing_dept
    int available
  }
  reservations {
    int id PK
    int user_id FK
    int venue_id FK
    int club_id FK
    datetime start_time
    datetime end_time
    string purpose
    string status
    int priority_order
    string ai_risk_level
    string ai_reasoning
    int reviewed_by FK
    datetime reviewed_at
    string pdf_url
    int conflict_reservation_id FK
  }
  conflicts {
    string id PK
    int venue_id FK
    date conflict_date
    int party1_user_id FK
    int party2_user_id FK
    string status
    datetime opened_at
    datetime ai_intervened_at
    datetime penalty_applied_at
    datetime resolved_at
    string resolution
    string ai_suggestion_text
  }
  credit_logs {
    int id PK
    int user_id FK
    string action
    int change_amount
    int score_after
    string reason
    string related_type
    int related_id
    datetime created_at
  }
  notification_logs {
    int id PK
    string token UK
    int user_id FK
    datetime read_at
    string ip_address
    string user_agent
    int is_mandatory
    int confirmed
  }
  map_elements {
    int id PK
    string building_id
    string building_name
    string status
    string maintenance_note
    datetime maintenance_end
    real lat
    real lng
    int accessible
    int has_elevator
    int has_restroom
    int is_reservable
  }
  equipment {
    int id PK
    string name
    string category
    int qty_total
    int qty_available
    string location
  }
  equipment_loans {
    int id PK
    int equipment_id FK
    int user_id FK
    date return_due_date
    string status
    string totp_secret
    datetime created_at
  }
  activities {
    int id PK
    int club_id FK
    string title
    date activity_date
    int venue_id FK
    string status
    string tags
  }
  portfolios {
    int id PK
    int user_id FK UK
    string bio
    string skills
    int volunteer_total_hours
  }
  portfolio_entries {
    int id PK
    int user_id FK
    string entry_type
    string title
    string role
    date date
    int hours
    string achievement
  }
  users ||--o{ reservations : "申請"
  venues ||--o{ reservations : "被預約"
  clubs ||--o{ reservations : "透過"
  reservations ||--o| conflicts : "產生"
  users ||--o{ credit_logs : "紀錄"
  users ||--o{ notification_logs : "已讀"
  users ||--o{ equipment_loans : "借用"
  equipment ||--o{ equipment_loans : "被借"
  clubs ||--o{ activities : "舉辦"
  venues ||--o{ activities : "場地"
  users ||--|| portfolios : "擁有"
  users ||--o{ portfolio_entries : "包含"
  users ||--o| conflicts : "party1"
  users ||--o| conflicts : "party2"
  venues ||--o| map_elements : "對應"
''')

    add_heading(doc, '4.2 資料表清單', level=2)
    tbl3 = doc.add_table(rows=14, cols=3)
    tbl3.style = 'Table Grid'
    add_table_header(tbl3, ['資料表', '用途', '關鍵欄位'])
    tables_data = [
        ('users', '使用者帳號', 'student_id, email, role, credit_score, two_fa_secret, jwt_version'),
        ('clubs', '社團資料', 'name, category, advisor_id, credit_score, status'),
        ('venues', '場地資料', 'name, capacity, type, managing_dept, available'),
        ('reservations', '場地預約（三階段）', 'user_id, venue_id, status, ai_risk_level, priority_order, weight_level'),
        ('conflicts', '衝突協商記錄', 'id, party1/2, status, opened_at, ai_intervened_at, penalty_applied_at'),
        ('credit_logs', '信用積分變動記錄', 'user_id, action, change_amount, score_after, reason'),
        ('notification_logs', '通知已讀回條', 'token, user_id, read_at, ip_address, is_mandatory, confirmed'),
        ('map_elements', '地圖建築動態狀態', 'building_id, status, maintenance_note, lat, lng, accessible'),
        ('equipment', '器材庫存', 'name, category, qty_total, qty_available, location'),
        ('equipment_loans', '器材借用記錄', 'equipment_id, user_id, return_due_date, status, totp_secret'),
        ('activities', '活動資料', 'club_id, title, activity_date, venue_id, status, tags'),
        ('portfolios', 'E-Portfolio 基本資料', 'user_id, bio, skills, volunteer_total_hours'),
        ('portfolio_entries', '履歷條目', 'user_id, entry_type, title, role, date, achievement'),
    ]
    for i, row_data in enumerate(tables_data):
        row = tbl3.rows[i+1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(8)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 5. 完整 API 規格
    # ════════════════════════════════════════════
    add_heading(doc, '5. 完整 API 規格', level=1)

    api_groups = [
        ('5.1 認證 API', [
            ('GET',  '/api/auth/me', '取得當前用戶資訊', 'Authorization: Bearer JWT'),
            ('POST', '/api/auth/google', '示範登入（Google OAuth）', 'Body: { code }'),
        ]),
        ('5.2 場地 API', [
            ('GET',  '/api/venues', '取得場地列表（含容量、類型、可用狀態）', ''),
            ('GET',  '/api/reservations', '取得預約列表', ''),
            ('POST', '/api/reservations', '新增預約申請（AI 預審）', 'Body: { venue_id, start_time, end_time, purpose, role }'),
            ('PATCH', '/api/reservations/:id/status', '更新預約狀態', 'Body: { status }'),
        ]),
        ('5.3 地圖 API', [
            ('GET', '/api/campus/buildings', '取得所有建築資訊（含無障礙欄位）', ''),
            ('GET', '/api/campus/accessibility', '取得無障礙設施彙總', ''),
            ('GET', '/api/map/building/:id', '動態獲取建築維護狀態（map_elements 表）', ''),
        ]),
        ('5.4 衝突協商 API', [
            ('GET',  '/api/conflicts', '取得衝突清單', ''),
            ('PATCH', '/api/conflicts/:id/resolve', '解決衝突', 'Body: { resolution, winner_party }'),
            ('POST', '/api/conflicts/:id/penalty', '施加罰款（扣 10 分）', 'Body: { user_id, reason }'),
        ]),
        ('5.5 信用積分 API', [
            ('GET',  '/api/credit/:userId', '取得用戶信用分數與歷史', ''),
            ('POST', '/api/credit/log', '記錄信用變動', 'Body: { user_id, action, change, score_after, reason }'),
            ('GET',  '/api/credit/logs/:userId', '取得信用變動記錄', ''),
        ]),
        ('5.6 通知與已讀追蹤 API', [
            ('POST', '/api/notifications/track', '記錄已讀回條（token + IP）', 'Body: { token, timestamp }'),
            ('GET',  '/api/notifications/logs', '取得已讀紀錄', ''),
        ]),
        ('5.7 AI 工具 API', [
            ('POST', '/api/ai/navigate', 'AI 校園導覽問答', 'Body: { message }'),
            ('POST', '/api/ai/regulations', '法規查詢 RAG', 'Body: { query }'),
            ('POST', '/api/ai/venue-workflow', '場地預約流程 RAG', 'Body: { requirement }'),
            ('POST', '/api/ai/screen', 'AI 預審申請', 'Body: { reason, purpose, role }'),
            ('POST', '/api/ai/generate-plan', 'AI 活動企劃生成', 'Body: { event_name, club_name, expected_participants, event_type }'),
        ]),
        ('5.8 其他 API', [
            ('GET', '/api/clubs', '社團列表', ''),
            ('GET', '/api/clubs/:id', '社團詳情', ''),
            ('GET', '/api/equipment', '器材庫存', ''),
            ('POST', '/api/equipment/borrow', '申請借用器材', 'Body: { equipment_id, return_date, user_id }'),
            ('GET', '/api/activities', '活動牆', 'Query: ?keyword=&tag='),
            ('GET', '/api/dashboard/:role', '儀表板數據', ''),
            ('GET', '/api/portfolio/:userId', 'E-Portfolio', ''),
            ('GET', '/api/calendar', '全域行事曆', 'Query: ?month=YYYY-MM'),
            ('GET', '/api/users', '用戶列表（管理員）', ''),
        ]),
    ]

    for group_title, endpoints in api_groups:
        add_heading(doc, group_title, level=2)
        tbl_api = doc.add_table(rows=len(endpoints)+1, cols=4)
        tbl_api.style = 'Table Grid'
        add_table_header(tbl_api, ['Method', 'Path', '說明', '參數'])
        for i, ep in enumerate(endpoints):
            row = tbl_api.rows[i+1]
            method_colors = {'GET':'E8F4FD', 'POST':'E8F8E8', 'PATCH':'FFF3E0', 'DELETE':'FFE0E0'}
            set_cell_bg(row.cells[0], method_colors.get(ep[0],'F1F5F9'))
            for j, val in enumerate(ep):
                row.cells[j].text = val
                for para in row.cells[j].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
                        if j == 0: run.font.bold = True
        doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 6. 安全性與認證架構
    # ════════════════════════════════════════════
    add_heading(doc, '6. 安全性與認證架構', level=1)

    add_heading(doc, '6.1 認證流程（Mermaid）', level=2)
    add_mermaid_block(doc, 'JWT + 2FA 認證流程', '''
sequenceDiagram
  participant U as 使用者
  participant FE as 前端
  participant BE as Hono Backend
  participant G as Google OAuth
  participant DB as D1 Database
  participant SMS as SMS/TOTP

  U->>FE: 點擊 Google 登入
  FE->>G: OAuth 重導向
  G-->>FE: Authorization Code
  FE->>BE: POST /api/auth/google { code }
  BE->>G: 驗證 code，取得用戶資訊
  G-->>BE: { email, name, google_id }
  BE->>BE: 驗證 @cloud.fju.edu.tw 域名
  BE->>DB: 查詢用戶 (email)
  alt 需要 2FA
    BE->>SMS: 發送 TOTP/SMS 驗證碼
    SMS-->>U: 驗證碼
    U->>FE: 輸入驗證碼
    FE->>BE: POST /api/auth/verify-2fa { code }
  end
  BE->>BE: 生成 JWT (含 role, credit, jwt_version)
  BE-->>FE: Set-Cookie: jwt=... HttpOnly Secure
  FE->>FE: 進入主應用
''')

    add_heading(doc, '6.2 CORS 設定', level=2)
    add_styled_para(doc, '允許來源：https://fju-smart-hub.pages.dev + http://localhost:3000')
    add_styled_para(doc, '允許方法：GET, POST, PUT, DELETE, PATCH')
    add_styled_para(doc, '允許標頭：Authorization, Content-Type')

    add_heading(doc, '6.3 JWT 失效機制（Observer Pattern）', level=2)
    add_styled_para(doc, '■ users.jwt_version 欄位記錄當前有效版本號')
    add_styled_para(doc, '■ 每次生成 JWT 時嵌入 jwt_version')
    add_styled_para(doc, '■ 驗證 JWT 時比對 DB 中的 jwt_version，不符則強制重登')
    add_styled_para(doc, '■ 管理員可遠端強制登出（increment jwt_version）')

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 7. 三階段預約流程
    # ════════════════════════════════════════════
    add_heading(doc, '7. 三階段場地預約流程', level=1)

    add_mermaid_block(doc, '三階段預約流程總覽', '''
flowchart LR
  A([社團申請人]) --> B[填寫申請表單<br/>Phase 1: 志願序申請]
  B --> C{AI 預審<br/>POST /api/ai/screen}
  C -->|High Risk| D[自動拒絕<br/>通知申請人]
  C -->|Medium Risk| E[人工審核<br/>行政人員介入]
  C -->|Low Risk| F{衝突偵測}
  E --> F
  F -->|無衝突| G[Phase 3: 官方核定<br/>發送 PDF 確認單]
  F -->|有衝突| H[Phase 2: 衝突協商<br/>LINE Notify 通知]
  H --> I{協商結果}
  I -->|協議達成| G
  I -->|3分鐘逾時| J[GPT-4 介入<br/>3項建議]
  J --> I
  I -->|6分鐘逾時| K[強制處理<br/>扣 10 信用分]
  G --> L[核准通知<br/>LINE + Outlook]
  G --> M[TOTP QR Code]
''')

    add_heading(doc, '7.1 志願序申請（Phase 1）', level=2)
    add_styled_para(doc, '■ 申請人填寫場地、時段、活動說明、志願序優先級（第1/2/3志願）')
    add_styled_para(doc, '■ AI 預審使用 Dify RAG，根據活動說明分析風險等級（Low/Medium/High）')
    add_styled_para(doc, '■ High Risk → 自動拒絕，附帶法規條文說明')
    add_styled_para(doc, '■ 預約記錄寫入 reservations 表，status = PENDING / AI_SCREENING')

    add_heading(doc, '7.2 衝突協商（Phase 2）', level=2)
    add_styled_para(doc, '■ 系統偵測到同時段同場地衝突時，建立 conflicts 記錄')
    add_styled_para(doc, '■ 雙方收到 LINE Notify 通知（含唯一 token，trigger read receipt）')
    add_styled_para(doc, '■ 協商室開啟，計時器啟動')
    add_styled_para(doc, '■ 3分鐘規則：沉默或語意無關 → GPT-4 介入，提供 3 個具體建議')
    add_styled_para(doc, '■ 6分鐘規則：無共識 → 強制關閉、扣 10 信用分、記錄 credit_logs')
    add_styled_para(doc, '■「協商完成」→ 更新 conflicts 表，加 2 信用分，跳轉行事曆管理')

    add_heading(doc, '7.3 官方核定（Phase 3）', level=2)
    add_styled_para(doc, '■ 行政人員（admin）登入後，在儀表板查看 pending_review 待審清單')
    add_styled_para(doc, '■ 核定後發送：LINE Notify + Microsoft Graph Outlook 行事曆邀請')
    add_styled_para(doc, '■ 核准通知附帶 TOTP QR Code，用於現場打卡確認')
    add_styled_para(doc, '■ reservations.status 更新為 APPROVED，pdf_url 生成確認單')

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 8. SWOT 分析
    # ════════════════════════════════════════════
    add_heading(doc, '8. SWOT 分析', level=1)
    tbl_swot = doc.add_table(rows=3, cols=2)
    tbl_swot.style = 'Table Grid'
    tbl_swot.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_swot = [('S 優勢（Strengths）', '003153'), ('W 劣勢（Weaknesses）', 'DAA520')]
    bodies = [
        '・AI GPT-4 自動介入協商，全球首創 3/6 分鐘仲裁機制\n・Leaflet.js 互動地圖 + 無障礙設施標記\n・玻璃態行事曆，視覺設計符合現代 UI 趨勢\n・JWT + 2FA + HttpOnly Cookie，安全性高\n・Cloudflare Edge 部署，全球延遲 < 50ms',
        '・示範版本使用靜態資料，實際 GPT-4 API 需要費用\n・Cloudflare Workers 10ms CPU 限制，複雜計算受限\n・Socket.io 即時通訊需要額外服務支援\n・LINE Notify 需要各用戶個別授權',
        '・整合 Microsoft Graph API，連接 Outlook 行事曆\n・擴展至全台大學，成為跨校園活動平台\n・AI 學習歷史協商資料，持續提升建議品質\n・電子 E-Portfolio 可對接人力銀行系統',
        '・OpenAI API 成本隨用量成長\n・GDPR/個資法合規要求\n・競爭產品（如 HiSKIO、Accupass）\n・學生習慣改變難度較高',
    ]
    for i, (header_text, color) in enumerate(headers_swot):
        cell = tbl_swot.rows[0].cells[i]
        cell.text = header_text
        set_cell_bg(cell, color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    for i in range(2):
        tbl_swot.rows[1].cells[i].text = bodies[i]
        tbl_swot.rows[1].cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    headers2 = [('O 機會（Opportunities）', '008000'), ('T 威脅（Threats）', 'FF0000')]
    for i, (header_text, color) in enumerate(headers2):
        cell = tbl_swot.rows[2].cells[i]
        cell.text = header_text
        set_cell_bg(cell, color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════
    # 9. 附錄
    # ════════════════════════════════════════════
    add_heading(doc, '9. 附錄', level=1)

    add_heading(doc, '9.1 版本變更記錄', level=2)
    tbl_change = doc.add_table(rows=5, cols=3)
    tbl_change.style = 'Table Grid'
    add_table_header(tbl_change, ['版本', '日期', '主要變更'])
    changes = [
        ('v1.0', '2025-09', '初始版本，基本預約功能'),
        ('v2.0', '2025-11', '新增 AI 預審、社團管理'),
        ('v3.1', '2026-03', '校園地圖、三大 AI 介面、無障礙設施'),
        ('v4.0', str(datetime.date.today()), '動畫引導頁、Leaflet 地圖、玻璃態行事曆、3/6分鐘協商、已讀回條、credit_logs、notification_logs、衝突 API'),
    ]
    for i, row_data in enumerate(changes):
        row = tbl_change.rows[i+1]
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    add_heading(doc, '9.2 名詞解釋', level=2)
    glossary = [
        ('Bento Grid', 'CSS Grid 佈局風格，以方格形式排列不同大小的資訊卡片'),
        ('Glassmorphism', '玻璃形態設計，使用 backdrop-filter:blur 實現半透明磨砂效果'),
        ('RAG', 'Retrieval-Augmented Generation，結合知識庫檢索的 AI 生成技術'),
        ('TOTP', 'Time-based One-Time Password，時間型一次性密碼（Google Authenticator）'),
        ('credit_logs', '信用積分變動記錄表，追蹤每次加扣分的原因、時間、操作者'),
        ('notification_logs', '通知已讀回條表，記錄用戶閱讀重要通知的時間戳與 IP'),
        ('map_elements', '地圖元素動態狀態表，側欄點擊時 AJAX 查詢建築維護狀態'),
        ('conflicts', '場地衝突協商記錄表，包含雙方、計時器觸發時間、處置結果'),
        ('flyTo()', 'Leaflet.js API，地圖平滑飛行至目標坐標，模擬 Google Maps 效果'),
        ('jwt_version', 'JWT 失效版本號（Observer Pattern），管理員可遠端強制登出用戶'),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.add_run(f'{term}：').bold = True
        p.add_run(definition)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, '9.3 部署環境', level=2)
    add_styled_para(doc, '■ 沙箱 Demo URL：https://3000-i12ven33yzeuey60x9hlh-2b54fc91.sandbox.novita.ai')
    add_styled_para(doc, '■ Cloudflare Pages（生產）：https://fju-smart-hub.pages.dev（待部署）')
    add_styled_para(doc, '■ GitHub Repository：https://github.com/KY0126/SA-v2-')
    add_styled_para(doc, '■ 技術棧：Hono v4 + TypeScript + Cloudflare Workers + D1 SQLite + KV + R2')

    # Save
    filename = '/home/user/webapp/FJU_Smart_Hub_SA文件_v4.0.docx'
    doc.save(filename)
    print(f'✅ Document saved: {filename}')
    import os
    print(f'   Size: {os.path.getsize(filename):,} bytes ({os.path.getsize(filename)/1024:.1f} KB)')
    return filename

if __name__ == '__main__':
    try:
        build_document()
    except Exception as e:
        print(f'❌ Error: {e}')
        import traceback; traceback.print_exc()
