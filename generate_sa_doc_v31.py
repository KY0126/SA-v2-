#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FJU Smart Hub v3.1 - System Analysis & Design Specification Generator
Generates Word document with auto Table of Contents and 5 system diagrams
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

def add_toc(doc):
    """Add automatic Table of Contents field"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)
    
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph

def set_cell_bg(cell, hex_color):
    """Set table cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    """Add heading with proper style"""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_para(doc, text, bold=False, size=11, indent=False):
    """Add formatted paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_info_box(doc, title, items, bg_color='F0F4FF'):
    """Add a styled info box table"""
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = 'Table Grid'
    
    # Header
    header_row = table.rows[0]
    set_cell_bg(header_row.cells[0], '003087')
    set_cell_bg(header_row.cells[1], '003087')
    header_row.cells[0].merge(header_row.cells[1])
    hdr = header_row.cells[0].paragraphs[0]
    hdr_run = hdr.add_run(title)
    hdr_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    hdr_run.font.bold = True
    hdr_run.font.size = Pt(11)
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content rows
    for i, (key, val) in enumerate(items):
        row = table.rows[i + 1]
        set_cell_bg(row.cells[0], 'E8EDF8')
        row.cells[0].text = key
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = val
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    # Set column widths
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(2.0 if idx == 0 else 4.5)
    
    doc.add_paragraph()
    return table

def add_diagram_placeholder(doc, title, description, color='003087'):
    """Add a diagram placeholder with description"""
    # Title box
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, color)
    title_para = cell.paragraphs[0]
    title_run = title_para.add_run(f'【圖示】{title}')
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    title_run.font.bold = True
    title_run.font.size = Pt(12)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Description box
    desc_table = doc.add_table(rows=1, cols=1)
    desc_table.style = 'Table Grid'
    desc_cell = desc_table.cell(0, 0)
    set_cell_bg(desc_cell, 'F8FAFC')
    desc_para = desc_cell.paragraphs[0]
    desc_run = desc_para.add_run(description)
    desc_run.font.size = Pt(10)
    desc_run.font.color.rgb = RGBColor(0x37, 0x47, 0x5A)
    desc_para.paragraph_format.space_before = Pt(6)
    desc_para.paragraph_format.space_after = Pt(6)
    desc_para.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph()
    return table

def generate_document():
    doc = Document()
    
    # ─── Page Setup ───
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # ─── Set Styles ───
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    
    # ─── Cover Page ───
    doc.add_paragraph()
    doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('FJU Smart Hub v3.1')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x00, 0x30, 0x87)
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('輔仁大學全方位校園管理系統')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(0x47, 0x48, 0x5C)
    
    doc.add_paragraph()
    
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2_run = sub2.add_run('系統分析與設計規格書')
    sub2_run.font.size = Pt(16)
    sub2_run.font.bold = True
    sub2_run.font.color.rgb = RGBColor(0xC8, 0xA9, 0x51)
    
    doc.add_paragraph()
    
    # Project info table on cover
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    cover_info = [
        ('專案名稱', 'FJU Smart Hub - 輔仁大學校園管理系統'),
        ('版本號', 'v3.1 (2026-04)'),
        ('文件類型', '系統分析與設計規格書（含 AI 模組擴充）'),
        ('作者/團隊', '資訊管理學系 系統分析課程小組'),
        ('審核者', '課外活動指導組 / 資訊中心'),
        ('最後更新', datetime.datetime.now().strftime('%Y/%m/%d')),
    ]
    for i, (k, v) in enumerate(cover_info):
        row = info_table.rows[i]
        set_cell_bg(row.cells[0], 'E8EDF8')
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = v
        if row.cells[1].paragraphs[0].runs:
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_page_break()
    
    # ─── TABLE OF CONTENTS ───
    add_heading(doc, '目錄 Table of Contents', level=1)
    
    toc_note = doc.add_paragraph()
    toc_note_run = toc_note.add_run('（請在 Microsoft Word 中按 Ctrl+A 全選後，再按 F9 鍵更新目錄）')
    toc_note_run.font.size = Pt(9)
    toc_note_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    toc_note_run.font.italic = True
    
    add_toc(doc)
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 1: SYSTEM OVERVIEW
    # ═══════════════════════════════════════════════
    add_heading(doc, '第一章  系統概述', level=1)
    
    add_heading(doc, '1.1 系統背景與目標', level=2)
    add_para(doc, '輔仁大學擁有 87 個學生社團及數千名學生，每學期產生大量場地預約、器材借用及活動審核需求。傳統人工處理流程效率低下，且缺乏透明的信用管理機制。')
    add_para(doc, 'FJU Smart Hub v3.1 整合 AI 智慧預審、RAG 法規查詢、三階段預約流程、無障礙校園地圖等功能，打造一個完整的校園資源管理生態系統。')
    
    add_info_box(doc, '系統核心目標', [
        ('效率提升', '將場地預約審核時間從 3-5 天縮短至 AI 自動審核 30 分鐘'),
        ('AI 整合', '透過 Dify AI 中間層實現自然語言查詢、RAG 法規搜尋'),
        ('無障礙服務', '整合無障礙設施地圖，保障身障學生的校園使用權'),
        ('信用管理', '全自動化信用點數追蹤，促進社團自律'),
        ('多語言支援', '支援繁中、英、日、韓、法五種語言介面'),
    ])
    
    add_heading(doc, '1.2 系統範圍', level=2)
    add_para(doc, '本系統涵蓋以下十大核心模組：')
    
    modules_table = doc.add_table(rows=11, cols=3)
    modules_table.style = 'Table Grid'
    # Header
    header = modules_table.rows[0]
    for cell in header.cells:
        set_cell_bg(cell, '003087')
    for text, cell in zip(['模組編號', '模組名稱', '主要功能'], header.cells):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    modules = [
        ('M01', 'E-Portfolio', '個人成長歷程、活動記錄、技能標籤'),
        ('M02', 'AI 企劃生成器', '自然語言輸入自動產生符合規範的企劃書'),
        ('M03', '幹部認證模組', 'TOTP QR Code 身份驗證、幹部認證管理'),
        ('M04', '器材庫存管理', '即時庫存、借用期限追蹤、逾期警示'),
        ('M05', 'AI 智慧預審', 'Dify 關鍵字風險評估、多層次審核機制'),
        ('M06', '校園地圖模組', '互動式地圖、無障礙設施標示、場地資訊'),
        ('M07', 'AI 申訴摘要', 'AI 自動生成申訴書草稿、申訴狀態追蹤'),
        ('M08', '動態活動牆', '社團活動展示、多標籤篩選、即時更新'),
        ('M09', '數位時光膠囊', '活動記憶保存、周年提醒'),
        ('M10', '全面二步驟驗證', 'PHPGangsta TOTP + Redis Session 管理'),
    ]
    for i, (num, name, func) in enumerate(modules):
        row = modules_table.rows[i + 1]
        if i % 2 == 0:
            set_cell_bg(row.cells[0], 'F8FAFC')
            set_cell_bg(row.cells[1], 'F8FAFC')
            set_cell_bg(row.cells[2], 'F8FAFC')
        row.cells[0].text = num
        row.cells[1].text = name
        row.cells[2].text = func
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 2: TECHNOLOGY STACK
    # ═══════════════════════════════════════════════
    add_heading(doc, '第二章  技術架構', level=1)
    
    add_heading(doc, '2.1 系統架構圖說明', level=2)
    add_para(doc, '本系統採用前後端分離架構，部署於 Cloudflare Pages/Workers 邊緣網路，透過 Dify AI 中間層整合多種 AI 能力。')
    
    add_diagram_placeholder(doc,
        '圖 2-1：系統架構圖（Dify AI 中間層 + Cloudflare WAF/Turnstile/R2 + 通知鏈）',
        '''系統架構說明（自上而下）：

【用戶端層】
  瀏覽器 (Vue 3 + GSAP) ← Cloudflare Turnstile (機器人防護)
  
【邊緣安全層】
  Cloudflare WAF (防 SQL Injection / XSS) → Rate Limiting → DDoS 防護
  
【應用層】
  Cloudflare Workers (Hono.js TypeScript) → API Routes
  ↓
【AI 中間層 (Dify)】
  ├── AI 導覽查詢 (LLM + Embedding)
  ├── 法規 RAG (Knowledge Base + Retrieval)
  ├── 預約流程 RAG (Workflow Agent)
  ├── AI 智慧預審 (Risk Classification)
  └── AI 申訴摘要 (Text Generation)
  ↓
【資料層】
  Cloudflare D1 (SQLite) + KV (Session) + R2 (Files)
  
【通知鏈】
  LINE Notify → SMTP (Outlook) → SMS (Twilio)
  觸發條件：預約狀態更新 / 信用點數警告 / 衝突協商邀請''',
        '003087'
    )
    
    add_heading(doc, '2.2 前端技術', level=2)
    add_info_box(doc, '前端技術棧', [
        ('框架', 'Vue 3 (CDN Style) + GSAP/ScrollTrigger 動畫'),
        ('國際化', 'vue-i18n 繁中/英/日/韓/法 五語言'),
        ('圖表', 'Chart.js 4.x (Line/Pie/Radar/Gauge/Heatmap)'),
        ('樣式', 'TailwindCSS + 自訂 CSS Variables'),
        ('地圖', 'SVG 互動式校園地圖 + 無障礙設施標示'),
        ('HTTP', 'Axios + 非同步 AI 任務輪詢'),
    ])
    
    add_heading(doc, '2.3 後端技術', level=2)
    add_info_box(doc, '後端技術棧（Cloudflare Workers）', [
        ('框架', 'Hono.js v4.x (TypeScript)'),
        ('部署', 'Cloudflare Pages + Workers (Edge Runtime)'),
        ('資料庫', 'Cloudflare D1 (SQLite) + Materialized Views'),
        ('快取/Session', 'Cloudflare KV + Redis (2FA Session)'),
        ('檔案儲存', 'Cloudflare R2 (S3-compatible)'),
        ('認證', 'Google OAuth (hd: cloud.fju.edu.tw) + PHPGangsta TOTP 2FA'),
        ('安全', 'Cloudflare WAF + Turnstile + JWT'),
        ('AI 串接', 'Dify API (非同步 Outbox Pattern)'),
    ])
    
    add_heading(doc, '2.4 資料庫設計', level=2)
    add_para(doc, '採用 Cloudflare D1 SQLite 資料庫，含 Materialized Views 加速統計查詢。')
    
    db_schemas = [
        ('users', '用戶表', 'id, name, student_id, email, phone, line_id, role, credit_score, totp_secret, created_at'),
        ('clubs', '社團表', 'id, name, category, advisor_id, member_count, credit_score, status, created_at'),
        ('venues', '場地表', 'id, name, capacity, type, accessible, elevator, restroom, reservable, location'),
        ('reservations', '預約表', 'id, user_id, venue_id, start_time, end_time, purpose, status, ai_risk_level, weight_score, pdf_path'),
        ('equipment', '器材表', 'id, name, category, qty_total, qty_available, location, condition'),
        ('borrow_records', '借用記錄表', 'id, user_id, equipment_id, qty, borrow_date, return_date, status, totp_verified'),
        ('credit_history', '信用記錄表', 'id, user_id, action, change, score_after, date, note'),
        ('ai_outbox', 'AI 非同步任務表', 'id, task_type, payload, status, result, created_at, processed_at'),
    ]
    
    schema_table = doc.add_table(rows=len(db_schemas)+1, cols=3)
    schema_table.style = 'Table Grid'
    header = schema_table.rows[0]
    for cell in header.cells:
        set_cell_bg(cell, '37474F')
    for text, cell in zip(['資料表', '用途', '主要欄位'], header.cells):
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
    
    for i, (tbl, desc, cols) in enumerate(db_schemas):
        row = schema_table.rows[i+1]
        if i % 2 == 0:
            for c in row.cells: set_cell_bg(c, 'F8FAFC')
        row.cells[0].text = tbl
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = desc
        row.cells[2].text = cols
        for c in row.cells[1:]:
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 3: CAMPUS MAP & ACCESSIBILITY
    # ═══════════════════════════════════════════════
    add_heading(doc, '第三章  校園地圖與無障礙設施模組', level=1)
    
    add_heading(doc, '3.1 功能概述', level=2)
    add_para(doc, '校園地圖模組提供互動式 SVG 地圖，標示所有主要建築、可預約場地及無障礙設施位置，協助師生快速找到所需資源。')
    
    add_info_box(doc, '地圖功能特色', [
        ('互動式 SVG 地圖', '點擊建築可查看詳細資訊、容量、無障礙設施'),
        ('無障礙設施標示', '♿ 坡道、🛗 電梯、🚻 無障礙廁所、P♿ 停車位'),
        ('篩選功能', '可篩選：全部/僅無障礙設施/可預約場地'),
        ('場地連結', '直接從地圖頁跳轉至場地預約或 AI 導覽'),
        ('多語言支援', '地圖說明文字支援 5 種語言'),
    ])
    
    add_heading(doc, '3.2 無障礙設施清單', level=2)
    
    access_table = doc.add_table(rows=9, cols=4)
    access_table.style = 'Table Grid'
    header = access_table.rows[0]
    for cell in header.cells:
        set_cell_bg(cell, '1565C0')
    for text, cell in zip(['建築名稱', '無障礙坡道', '無障礙電梯', '無障礙廁所'], header.cells):
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buildings_access = [
        ('焯炤館演講廳', '✅', '❌', '✅'),
        ('仁愛空間', '✅', '❌', '✅'),
        ('進修部演講廳', '✅', '✅', '❌'),
        ('圖書館', '✅', '✅（B1-6F）', '✅'),
        ('行政大樓', '✅', '✅（1-5F）', '✅'),
        ('學生活動中心', '✅', '✅（1-3F）', '✅'),
        ('宗倫樓', '✅', '✅（1-4F）', '✅'),
        ('體育館', '✅', '❌', '✅'),
    ]
    for i, row_data in enumerate(buildings_access):
        row = access_table.rows[i+1]
        if i % 2 == 0:
            for c in row.cells: set_cell_bg(c, 'F8FAFC')
        for j, text in enumerate(row_data):
            row.cells[j].text = text
            if row.cells[j].paragraphs[0].runs:
                row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
                if j > 0:
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 4: AI MODULES
    # ═══════════════════════════════════════════════
    add_heading(doc, '第四章  AI 人工智慧模組', level=1)
    
    add_heading(doc, '4.1 AI 導覽查詢介面', level=2)
    add_para(doc, 'AI 導覽查詢介面提供 24/7 即時問答服務，使用者可以自然語言詢問校園場地位置、預約流程、信用點數制度等問題。')
    
    add_info_box(doc, 'AI 導覽查詢規格', [
        ('介面類型', '對話式聊天介面（Chat-style）'),
        ('AI 後端', 'Dify LLM + FJU 校園知識庫'),
        ('知識庫來源', 'activity.fju.edu.tw / FB: submarinesky / 預約系統'),
        ('回覆能力', '場地位置、預約流程、信用點數、無障礙設施、器材借用'),
        ('快捷問題', '6 個 FAQ 快捷鍵，一鍵提問'),
        ('對話記憶', '單次 Session 內保留對話歷史'),
        ('語言支援', '繁中、英、日、韓、法'),
    ])
    
    add_heading(doc, '4.2 法規查詢 RAG', level=2)
    add_para(doc, '法規查詢 RAG 利用 Retrieval Augmented Generation 技術，從結構化法規知識庫中精準檢索相關條文，並附上引用來源。')
    
    add_info_box(doc, '法規 RAG 規格', [
        ('技術實作', 'Dify Knowledge Base + Embedding 向量搜尋'),
        ('知識庫範圍', '輔大學生社團活動管理辦法、場地管理辦法、器材借用辦法'),
        ('查詢介面', '自由文字輸入 + 8 個常見問題快捷'),
        ('回覆品質', '含信心度評分（0-100%）+ 引用條文來源'),
        ('更新頻率', '每學期更新一次知識庫'),
        ('API 端點', 'POST /api/ai/regulations'),
    ])
    
    add_heading(doc, '4.3 場地與器材預約流程 RAG', level=2)
    add_para(doc, '預約流程 RAG 根據使用者的預約需求，自動規劃最佳步驟路徑，並提供針對性的注意事項。')
    
    add_info_box(doc, '預約流程 RAG 規格', [
        ('技術實作', 'Dify Workflow Agent + 場景分類'),
        ('支援情境', '場地預約、器材借用、衝突處理、申訴流程'),
        ('輸出格式', '步驟清單 + 注意事項 + 快速操作按鈕'),
        ('器材清單', '即時連結器材庫存 API'),
        ('API 端點', 'POST /api/ai/venue-workflow'),
    ])
    
    add_heading(doc, '4.4 Dify AI 工作流程圖', level=2)
    add_diagram_placeholder(doc,
        '圖 4-1：Dify AI 與 RAG 工作流程圖',
        '''Dify AI 中間層工作流程：

【輸入層】
  用戶自然語言查詢 → API Gateway (Hono.js)
  ↓
【Dify Workflow Engine】
  1. 意圖識別（Intent Classification）
     ├── 場地導覽類 → AI 導覽 Agent
     ├── 法規查詢類 → Regulations RAG Chain
     ├── 預約流程類 → Venue Workflow Agent  
     ├── 風險評估類 → AI 預審 Classifier
     └── 申訴類 → Appeal Summary Generator
  ↓
  2. 知識庫檢索（Knowledge Base Retrieval）
     - Embedding Model: text-embedding-ada-002
     - Vector Store: Dify 內建向量資料庫
     - Top-K: 5 個最相關文件片段
  ↓
  3. LLM 生成（Response Generation）
     - 結合檢索結果 + 使用者查詢
     - 引用法規來源標注
     - 信心度計算
  ↓
【輸出層】
  結構化 JSON 回應 → 前端渲染
  非同步重任務 → Outbox Table → Background Worker''',
        '7C3AED'
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 5: RESERVATION WORKFLOW
    # ═══════════════════════════════════════════════
    add_heading(doc, '第五章  三階段預約工作流程', level=1)
    
    add_heading(doc, '5.1 預約狀態機說明', level=2)
    add_para(doc, '預約申請採用三階段狀態機設計，每個狀態轉換均有明確的觸發條件與通知機制。')
    
    add_diagram_placeholder(doc,
        '圖 5-1：三步驟預約狀態機（Reservation State Machine）',
        '''預約狀態機（3 States × 7 Transitions）：

初始狀態: DRAFT（草稿）
↓ [用戶提交]
─────────────────────────────────────────
第一階段：AI 智慧預審
─────────────────────────────────────────
PENDING_AI_REVIEW
  ├─[AI Risk=Low]──→ AI_APPROVED（自動核准）
  ├─[AI Risk=Medium]→ PENDING_MANUAL_REVIEW（待人工審核）
  └─[AI Risk=High]─→ AI_REJECTED（自動駁回）

─────────────────────────────────────────
第二階段：衝突協商
─────────────────────────────────────────
CONFLICT_DETECTED（偵測到時段衝突）
  ├─[協商成功/48h內]→ CONFLICT_RESOLVED
  └─[逾時未解決]───→ PENDING_ADMIN_ARBITRATION（裁定）

─────────────────────────────────────────
第三階段：官方核定
─────────────────────────────────────────
PENDING_OFFICIAL_APPROVAL
  ├─[課指組核准]──→ OFFICIALLY_APPROVED（發送 PDF + TOTP QR）
  └─[課指組駁回]──→ OFFICIALLY_REJECTED（通知理由）

─────────────────────────────────────────
信用點數聯動：
  APPROVED → 準時出席 → +2 分
  APPROVED → 未出席 → -10 分
  APPROVED → 臨時取消(24h內) → -3 分
  器材逾期 → -5 分/天''',
        'E65100'
    )
    
    add_heading(doc, '5.2 角色分工泳道圖', level=2)
    add_diagram_placeholder(doc,
        '圖 5-2：三階段預約角色分工泳道圖（Role-based Swimlane）',
        '''預約流程泳道分工（4 個角色）：

┌────────────┬──────────────────────────────────────────────────────┐
│  社團幹部   │ 填寫申請表 → AI預審 → 接收結果通知 → 衝突協商 → 簽到 │
├────────────┼──────────────────────────────────────────────────────┤
│  AI 系統   │ 關鍵字分析 → 風險評分 → 衝突偵測 → 自動通知 → 記錄   │
├────────────┼──────────────────────────────────────────────────────┤
│  課指組    │ 人工審核 → 裁定衝突 → 官方核定 → 發送PDF確認單        │
├────────────┼──────────────────────────────────────────────────────┤
│  指導教授  │ 審核活動企劃 → 簽核同意 → 接收通知                    │
└────────────┴──────────────────────────────────────────────────────┘

通知方式（依事件觸發）：
  申請提交    → Email (SMTP/Outlook)
  AI 預審完成 → 系統通知 + LINE Notify
  衝突偵測    → LINE Notify（所有相關社團）
  官方核定    → Email PDF + 系統通知 + SMS''',
        '37474F'
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 6: SECURITY & AUTH
    # ═══════════════════════════════════════════════
    add_heading(doc, '第六章  安全機制與身份驗證', level=1)
    
    add_heading(doc, '6.1 安全架構概述', level=2)
    add_diagram_placeholder(doc,
        '圖 6-1：安全與認證生命週期圖（Security & Auth Lifecycle）',
        '''完整認證流程（多層安全防護）：

【第一層：邊緣防護】
  Cloudflare WAF → DDoS 防護 → Rate Limiting (100 req/min)
  Cloudflare Turnstile → 機器人偵測（登入頁面前置）

【第二層：身份認證】
  Google OAuth 2.0 → hd 參數檢查（僅允許 cloud.fju.edu.tw）
  ↓
  JWT Token 簽發（有效期 8 小時）
  ↓
  TOTP 2FA 驗證（PHPGangsta/Google Authenticator）
  Secret 儲存：users.totp_secret（AES-256 加密）
  
【第三層：Session 管理】
  Redis 儲存 2FA Session（TTL: 30 分鐘）
  JWT Refresh Token（有效期 30 天）
  裝置指紋綁定（User-Agent + IP）

【第四層：API 授權】
  Role-Based Access Control (RBAC)
  ├── student: 讀取公開資源 + 個人申請
  ├── club_officer: 社團申請 + 器材借用
  ├── professor: 社團審核
  ├── admin: 全部功能 + 管理後台
  └── it_admin: 系統監控 + 用戶管理

【第五層：資料保護】
  API Key / Secrets → Cloudflare Secrets（環境變數）
  敏感資料加密儲存（AES-256）
  SQL Injection 防護（Parameterized Queries）
  XSS 防護（Content Security Policy）''',
        'B71C1C'
    )
    
    add_heading(doc, '6.2 安全設定規範', level=2)
    add_info_box(doc, '安全設定要求', [
        ('API Key 管理', '所有 API Key 存於 Cloudflare Secrets，禁止 hardcode'),
        ('Google OAuth', 'hd 參數強制檢查 cloud.fju.edu.tw 域名'),
        ('2FA 強制', '課指組職員、社團幹部必須啟用 TOTP 2FA'),
        ('密碼政策', '最少 8 碼，含大小寫英文及數字'),
        ('CORS 設定', '僅允許 fju-smart-hub.pages.dev 跨域請求'),
        ('資料庫', 'D1 SQL 使用 Parameterized Statements 防注入'),
        ('日誌審計', '所有管理操作記錄至 audit_log 表'),
    ])
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 7: USER STORIES & REQUIREMENTS
    # ═══════════════════════════════════════════════
    add_heading(doc, '第七章  使用者故事與需求規格', level=1)
    
    add_heading(doc, '7.1 核心使用者故事', level=2)
    
    user_stories = [
        ('US-01', '學生', '使用 AI 導覽助手查詢場地', '能快速找到需要的場地和無障礙設施', '高'),
        ('US-02', '社團幹部', '申請三階段場地預約', '透明化的審核進度和即時通知', '高'),
        ('US-03', '社團幹部', '使用 AI 企劃生成器', '快速產出合規的活動企劃書', '高'),
        ('US-04', '課指組', '使用 AI 法規 RAG 查詢', '快速找到相關法規條文', '中'),
        ('US-05', '學生', '查詢個人信用點數', '了解點數狀態和歷史記錄', '中'),
        ('US-06', '指導教授', '審核社團活動申請', '有效率地審核並給予意見', '中'),
        ('US-07', '學生', '使用無障礙校園地圖', '找到有無障礙設施的建築和路線', '高'),
        ('US-08', '社團幹部', '申訴被拒絕的預約', 'AI 協助生成申訴書草稿', '低'),
    ]
    
    us_table = doc.add_table(rows=len(user_stories)+1, cols=5)
    us_table.style = 'Table Grid'
    header = us_table.rows[0]
    for cell in header.cells:
        set_cell_bg(cell, '003087')
    for text, cell in zip(['編號', '角色', '故事', '價值', '優先'], header.cells):
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
    
    for i, row_data in enumerate(user_stories):
        row = us_table.rows[i+1]
        if i % 2 == 0:
            for c in row.cells: set_cell_bg(c, 'F8FAFC')
        for j, text in enumerate(row_data):
            row.cells[j].text = text
            if row.cells[j].paragraphs[0].runs:
                row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
                if row_data[4] == '高' and j == 4:
                    row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x30, 0x87)
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 8: API SPECIFICATION
    # ═══════════════════════════════════════════════
    add_heading(doc, '第八章  API 規格文件', level=1)
    
    add_heading(doc, '8.1 API 端點總覽', level=2)
    
    api_endpoints = [
        ('GET', '/api/venues', '取得場地列表', '所有角色'),
        ('POST', '/api/reservations', '建立預約申請', '已登入用戶'),
        ('GET', '/api/reservations', '查詢預約記錄', '已登入用戶'),
        ('PATCH', '/api/reservations/:id/status', '更新預約狀態', '課指組'),
        ('GET', '/api/equipment', '查詢器材庫存', '所有角色'),
        ('POST', '/api/equipment/borrow', '申請器材借用', '已登入用戶'),
        ('GET', '/api/clubs', '取得社團列表', '所有角色'),
        ('GET', '/api/dashboard/:role', '取得角色儀表板資料', '已登入用戶'),
        ('POST', '/api/ai/navigate', 'AI 導覽查詢', '已登入用戶'),
        ('POST', '/api/ai/regulations', '法規 RAG 查詢', '已登入用戶'),
        ('POST', '/api/ai/venue-workflow', '預約流程 RAG', '已登入用戶'),
        ('POST', '/api/ai/screen', 'AI 預審申請', '已登入用戶'),
        ('POST', '/api/ai/generate-plan', 'AI 企劃生成', '社團幹部'),
        ('GET', '/api/campus/buildings', '校園建築資訊', '所有角色'),
        ('GET', '/api/campus/accessibility', '無障礙設施查詢', '所有角色'),
        ('GET', '/api/credit/:userId', '查詢信用點數', '本人或管理員'),
    ]
    
    api_table = doc.add_table(rows=len(api_endpoints)+1, cols=4)
    api_table.style = 'Table Grid'
    header = api_table.rows[0]
    for cell in header.cells:
        set_cell_bg(cell, '37474F')
    for text, cell in zip(['Method', 'Endpoint', '功能說明', '權限'], header.cells):
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
    
    method_colors = {'GET': 'E3F2FD', 'POST': 'E8F5E9', 'PATCH': 'FFF3E0', 'DELETE': 'FFEBEE'}
    for i, (method, endpoint, desc, perm) in enumerate(api_endpoints):
        row = api_table.rows[i+1]
        bg = method_colors.get(method, 'F8FAFC')
        for c in row.cells: set_cell_bg(c, bg)
        row.cells[0].text = method
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = endpoint
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[2].text = desc
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[3].text = perm
        row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # CHAPTER 9: DEPLOYMENT & OPERATIONS
    # ═══════════════════════════════════════════════
    add_heading(doc, '第九章  部署與維運', level=1)
    
    add_heading(doc, '9.1 部署環境', level=2)
    add_info_box(doc, '部署架構', [
        ('平台', 'Cloudflare Pages + Workers（邊緣部署）'),
        ('前端 Demo URL', 'https://fju-smart-hub.pages.dev'),
        ('GitHub 倉庫', 'https://github.com/KY0126/SA-v2-'),
        ('CI/CD', 'GitHub Actions → Cloudflare Pages 自動部署'),
        ('環境', 'Production (main branch) + Preview (PR branches)'),
        ('監控', 'Cloudflare Analytics + Workers Metrics'),
        ('備份', 'D1 自動快照 + R2 備份策略'),
    ])
    
    add_heading(doc, '9.2 知識庫資料來源', level=2)
    add_info_box(doc, 'Dify RAG 知識庫來源', [
        ('法規文件', 'activity.fju.edu.tw（政策頁面、表單、社團清單）'),
        ('場地系統', 'http://140.136.202.67/fjuactivity/ + fjuspace/'),
        ('社群媒體', 'https://www.facebook.com/submarinesky/'),
        ('更新頻率', '法規文件：每學期；場地資料：每週'),
        ('索引方式', 'Dify Embedding + 自動切分（chunk_size: 500）'),
    ])
    
    add_heading(doc, '9.3 SWOT 分析', level=2)
    
    swot_table = doc.add_table(rows=3, cols=2)
    swot_table.style = 'Table Grid'
    swot_data = [
        ('優勢 (Strengths)', 
         '✅ AI 全自動審核大幅降低人力\n✅ Cloudflare 邊緣部署低延遲\n✅ 多語言支援服務國際學生\n✅ 無障礙設施整合關懷特殊需求',
         '劣勢 (Weaknesses)',
         '⚠️ Cloudflare Workers 10ms CPU 限制\n⚠️ 依賴 Dify AI 服務可用性\n⚠️ D1 SQLite 不支援複雜 JOIN'),
        ('機會 (Opportunities)',
         '🚀 AI 技術持續進步，RAG 準確度提升\n🚀 擴展至其他大學校園\n🚀 整合 LINE 官方帳號',
         '威脅 (Threats)',
         '⛔ AI 幻覺問題影響法規準確性\n⛔ 個資保護法規限制\n⛔ 資訊系統安全攻擊風險'),
    ]
    
    colors = [('D1FAE5', 'FEE2E2'), ('DBEAFE', 'FEF3C7')]
    titles = [('優勢 Strengths', '劣勢 Weaknesses'), ('機會 Opportunities', '威脅 Threats')]
    title_colors = ['059669', 'DC2626', '1D4ED8', 'D97706']
    
    for row_idx in range(2):
        for col_idx in range(2):
            cell = swot_table.rows[row_idx + 1].cells[col_idx]
            set_cell_bg(cell, colors[row_idx][col_idx])
            p = cell.paragraphs[0]
            run = p.add_run(titles[row_idx][col_idx] + '\n')
            run.font.bold = True
            run.font.size = Pt(11)
            content_run = p.add_run(swot_data[row_idx][col_idx*2 + 1])
            content_run.font.size = Pt(10)
    
    # Merge header
    header_row = swot_table.rows[0]
    header_row.cells[0].merge(header_row.cells[1])
    set_cell_bg(header_row.cells[0], '003087')
    header_para = header_row.cells[0].paragraphs[0]
    header_run = header_para.add_run('SWOT 分析 | FJU Smart Hub v3.1')
    header_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    header_run.font.bold = True
    header_run.font.size = Pt(12)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════
    # APPENDIX
    # ═══════════════════════════════════════════════
    add_heading(doc, '附錄 A：修改記錄', level=1)
    
    changelog_table = doc.add_table(rows=5, cols=4)
    changelog_table.style = 'Table Grid'
    header = changelog_table.rows[0]
    for cell in header.cells:
        set_cell_bg(cell, '003087')
    for text, cell in zip(['版本', '日期', '修改內容', '負責人'], header.cells):
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
    
    changes = [
        ('v1.0', '2026/01/15', '初始版本，基礎場地預約系統', 'SA 課程小組'),
        ('v2.0', '2026/02/20', '新增 AI 預審、信用點數系統', 'SA 課程小組'),
        ('v3.0', '2026/03/31', '新增儀表板、E-Portfolio、多語言', 'SA 課程小組'),
        ('v3.1', '2026/04/01', '新增校園地圖、3個 AI 介面（導覽/法規RAG/預約RAG）', 'SA 課程小組'),
    ]
    for i, row_data in enumerate(changes):
        row = changelog_table.rows[i+1]
        if i == len(changes) - 1:
            for c in row.cells: set_cell_bg(c, 'FFF3E0')
        for j, text in enumerate(row_data):
            row.cells[j].text = text
            if row.cells[j].paragraphs[0].runs:
                row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
                if i == len(changes) - 1:
                    row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    add_heading(doc, '附錄 B：名詞解釋', level=1)
    
    terms = [
        ('RAG', 'Retrieval-Augmented Generation - 檢索增強生成，結合知識庫搜尋與語言模型回答'),
        ('Dify', '開源 AI 應用開發平台，提供 LLM 編排、知識庫管理等功能'),
        ('TOTP', 'Time-based One-Time Password - 基於時間的一次性密碼，用於雙因素驗證'),
        ('Cloudflare WAF', 'Web Application Firewall - 網頁應用防火牆，防護 SQL 注入、XSS 等攻擊'),
        ('D1', 'Cloudflare 的 SQLite 相容資料庫服務，支援邊緣部署'),
        ('Hono.js', '輕量級 TypeScript 網頁框架，專為 Cloudflare Workers 設計'),
        ('Materialized View', '物化視圖，預先計算並儲存複雜查詢結果，提升查詢效能'),
        ('信用點數', '本系統設計的社團自律機制，初始100分，依行為加扣分'),
    ]
    
    terms_table = doc.add_table(rows=len(terms)+1, cols=2)
    terms_table.style = 'Table Grid'
    header = terms_table.rows[0]
    for cell in header.cells:
        set_cell_bg(cell, '37474F')
    for text, cell in zip(['術語', '說明'], header.cells):
        run = cell.paragraphs[0].add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
    
    for i, (term, desc) in enumerate(terms):
        row = terms_table.rows[i+1]
        if i % 2 == 0:
            set_cell_bg(row.cells[0], 'F8FAFC')
            set_cell_bg(row.cells[1], 'F8FAFC')
        row.cells[0].text = term
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = desc
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Footer note
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f'FJU Smart Hub v3.1 系統分析與設計規格書 | '
        f'輔仁大學資訊管理學系 | '
        f'生成日期：{datetime.datetime.now().strftime("%Y/%m/%d %H:%M")}'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    footer_run.font.italic = True
    
    # Save
    output_path = '/home/user/webapp/FJU_Smart_Hub_SA文件_v3.1.docx'
    doc.save(output_path)
    print(f'✅ 文件已生成：{output_path}')
    
    import os
    size = os.path.getsize(output_path)
    print(f'   文件大小：{size:,} bytes ({size/1024:.1f} KB)')
    return output_path

if __name__ == '__main__':
    generate_document()
